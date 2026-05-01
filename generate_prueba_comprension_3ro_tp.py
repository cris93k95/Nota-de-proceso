#!/usr/bin/env python3
"""Genera pruebas de comprension lectora TP para los cinco 3ros medios."""

import base64
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
PLAN_OUTPUT_DIR = ROOT / "PLANIFICACIONES_2026_LISTO_IMPRESION"
SOURCE_INSTRUMENTS_DIR = ROOT / "materiales-clases" / "3ro-medio" / "instrumentos"
PUBLISHED_INSTRUMENTS_DIR = ROOT / "tranquiprofe.cl" / "static" / "recursos" / "materiales" / "3ro-medio" / "instrumentos"
TARGET_DIRS = [SOURCE_INSTRUMENTS_DIR, PUBLISHED_INSTRUMENTS_DIR]
LOGO_PATH = ROOT / "_logo_header_resized.png"
LOGO_B64 = base64.b64encode(LOGO_PATH.read_bytes()).decode("ascii") if LOGO_PATH.exists() else ""
TOTAL_POINTS = 30
PASSING_POINTS = int(TOTAL_POINTS * 0.6)


COURSE_ASSESSMENTS = [
    {
        "file_stub": "PRUEBA_COMPRENSION_LECTORA_3RO_INDUSTRIAL",
        "course": "3°A — Mecánica Industrial",
        "course_field": "3°A ____",
        "specialty": "Mecánica Industrial",
        "objective": "Evaluar la comprensión de textos técnicos vinculados a procesos productivos, mantenimiento y seguridad en Mecánica Industrial.",
        "skills": "Comprensión lectora — vocabulario técnico — interpretación de procedimientos",
        "parts": [
            {
                "title": "PART I — Predictive Maintenance",
                "instructions": "Questions 1–5. Read the text and choose the correct answer. Each correct answer is worth 2 points.",
                "text_title": "Text 1 — Predictive Maintenance in a Packaging Plant",
                "text": (
                    "At a packaging plant in Rancagua, the industrial maintenance team noticed that one conveyor motor was stopping "
                    "for a few minutes every afternoon. Instead of waiting for the motor to fail, the technicians checked the vibration "
                    "sensor and the temperature record stored in the control panel. The data showed that the motor was getting hotter "
                    "after long production runs. The team cleaned the ventilation grid, adjusted the belt tension, and scheduled a short "
                    "maintenance stop before the next shift. After that intervention, the conveyor worked more smoothly and the line lost "
                    "less time. The supervisor explained that predictive maintenance is useful because it solves small problems before they "
                    "become expensive breakdowns."
                ),
                "questions": [
                    {
                        "num": 1,
                        "q": "What problem did the team detect at the beginning of the text?",
                        "options": [
                            "A) The control panel stopped working completely.",
                            "B) The conveyor motor stopped for short periods every afternoon.",
                            "C) The packaging line had no vibration sensor installed.",
                            "D) The supervisor could not read the maintenance report.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 2,
                        "q": "Which sources of information did the technicians review?",
                        "options": [
                            "A) The production budget and the operator schedule.",
                            "B) The conveyor manual and the shift calendar.",
                            "C) The vibration sensor and the temperature record.",
                            "D) The belt invoice and the emergency log.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 3,
                        "q": "Why did the team clean the ventilation grid and adjust the belt tension?",
                        "options": [
                            "A) Because the data showed the motor was overheating after long runs.",
                            "B) Because a new motor arrived from the supplier.",
                            "C) Because the supervisor wanted to change the whole line.",
                            "D) Because the conveyor was too noisy in the morning only.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 4,
                        "q": "What happened after the intervention?",
                        "options": [
                            "A) The line had to stop for a full day.",
                            "B) The motor needed to be replaced immediately.",
                            "C) The conveyor ran more smoothly and the line lost less time.",
                            "D) The team removed the vibration sensor from the plant.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 5,
                        "q": "What idea does the supervisor emphasize?",
                        "options": [
                            "A) Predictive maintenance avoids reading technical data.",
                            "B) Predictive maintenance helps solve small issues before they become costly failures.",
                            "C) Predictive maintenance is only useful during night shifts.",
                            "D) Predictive maintenance replaces all manual inspection.",
                        ],
                        "answer": "B",
                    },
                ],
            },
            {
                "title": "PART II — Procedure Sheet",
                "instructions": "Questions 6–10. Read the workshop procedure carefully and identify the key technical details.",
                "text_title": "Text 2 — Shaft Alignment Procedure",
                "text": (
                    "Before aligning a motor and a pump, the students in the workshop read the procedure sheet carefully. First, they "
                    "disconnect the electrical supply and place a lockout tag on the panel. Next, they clean the coupling surfaces and "
                    "install the dial indicator. Then they rotate the shaft slowly and record the measurements at four positions. If the "
                    "readings show a difference, the team adds or removes thin metal shims under the motor base. Finally, they tighten the "
                    "bolts again and repeat the measurements to confirm the alignment. The teacher reminds them that a dirty surface or a "
                    "missing measurement can produce vibration and damage the machine."
                ),
                "questions": [
                    {
                        "num": 6,
                        "q": "What is the first action in the alignment procedure?",
                        "options": [
                            "A) Rotate the shaft and record four measurements.",
                            "B) Tighten the bolts and test the pump.",
                            "C) Disconnect the electrical supply and place a lockout tag.",
                            "D) Install the new coupling on the machine.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 7,
                        "q": "Which tool is installed to check the shaft position?",
                        "options": [
                            "A) A wrench.",
                            "B) A dial indicator.",
                            "C) A soldering iron.",
                            "D) A hydraulic jack.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 8,
                        "q": "Why are shims added or removed under the motor base?",
                        "options": [
                            "A) To change the pump speed during production.",
                            "B) To correct the difference shown in the measurements.",
                            "C) To clean the shaft before testing.",
                            "D) To lock the electrical panel more securely.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 9,
                        "q": "What do the students do after tightening the bolts again?",
                        "options": [
                            "A) They repeat the measurements to confirm alignment.",
                            "B) They remove the dial indicator permanently.",
                            "C) They start the machine without supervision.",
                            "D) They paint the motor base blue.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 10,
                        "q": "According to the teacher, what can happen if the surface is dirty or a measurement is missing?",
                        "options": [
                            "A) The panel will show a battery warning light.",
                            "B) The workshop will lose the procedure sheet.",
                            "C) Vibration can appear and damage the machine.",
                            "D) The pump will produce colder water.",
                        ],
                        "answer": "C",
                    },
                ],
            },
            {
                "title": "PART III — Safety Notice",
                "instructions": "Questions 11–15. Read the safety notice and identify precautions, responsibilities, and the main risk being prevented.",
                "text_title": "Text 3 — Workshop Safety Notice",
                "text": (
                    "NOTICE: When moving a heavy motor housing, inspect the lifting slings before use. Keep the load low while transporting "
                    "it across the workshop. Only the assigned spotter may give hand signals to the operator. If oil is found on the floor, "
                    "stop the movement and clean the area before continuing. Never place your hands under the suspended load, even for a moment."
                ),
                "questions": [
                    {
                        "num": 11,
                        "q": "What must be inspected before moving the load?",
                        "options": [
                            "A) The work uniforms.",
                            "B) The lifting slings.",
                            "C) The classroom projector.",
                            "D) The maintenance calendar.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 12,
                        "q": "How should the load be transported across the workshop?",
                        "options": [
                            "A) As high as possible to avoid obstacles.",
                            "B) With two operators pulling from both sides.",
                            "C) Low to the ground while it is being moved.",
                            "D) Only after the motor is disassembled.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 13,
                        "q": "Who is allowed to give hand signals to the operator?",
                        "options": [
                            "A) Any student near the machine.",
                            "B) The assigned spotter only.",
                            "C) The person cleaning the floor.",
                            "D) The last student who used the crane.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 14,
                        "q": "What must happen if oil is found on the floor?",
                        "options": [
                            "A) The load should move faster to avoid delay.",
                            "B) The spotter should change places with the operator.",
                            "C) The movement must stop and the area must be cleaned.",
                            "D) The slings should be replaced immediately.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 15,
                        "q": "What is the main safety purpose of this notice?",
                        "options": [
                            "A) To explain how to paint the workshop floor.",
                            "B) To prevent crush injuries while moving suspended loads.",
                            "C) To schedule predictive maintenance after each shift.",
                            "D) To reduce the price of replacement motors.",
                        ],
                        "answer": "B",
                    },
                ],
            },
        ],
    },
    {
        "file_stub": "PRUEBA_COMPRENSION_LECTORA_3RO_AUTOMOTRIZ",
        "course": "3°B — Mecánica Automotriz",
        "course_field": "3°B ____",
        "specialty": "Mecánica Automotriz",
        "objective": "Evaluar la comprensión de textos técnicos vinculados a diagnóstico, mantenimiento y seguridad en Mecánica Automotriz.",
        "skills": "Comprensión lectora — vocabulario técnico — interpretación de protocolos",
        "parts": [
            {
                "title": "PART I — Diagnostic Report",
                "instructions": "Questions 1–5. Read the report and choose the best answer. Each correct answer is worth 2 points.",
                "text_title": "Text 1 — Hybrid Vehicle Diagnostic Report",
                "text": (
                    "During a service class, a group received a hybrid car with a warning light on the dashboard. The students connected a "
                    "scan tool and read a fault code related to battery temperature. At first, they thought the battery pack was damaged, but "
                    "the teacher asked them to inspect the cooling system first. They discovered that the air filter for the battery fan was "
                    "full of dust, so the fan could not move enough air. After cleaning the filter and clearing the code, they tested the "
                    "vehicle again. The warning light did not return, and the group concluded that a simple blockage can sometimes create a "
                    "serious diagnostic signal."
                ),
                "questions": [
                    {
                        "num": 1,
                        "q": "What problem did the group observe first?",
                        "options": [
                            "A) The brake pedal was too soft.",
                            "B) A warning light appeared on the dashboard.",
                            "C) The engine oil level was too low.",
                            "D) The tyres had uneven pressure.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 2,
                        "q": "Which tool did the students use to read the fault code?",
                        "options": [
                            "A) A scan tool.",
                            "B) A torque wrench.",
                            "C) A hydraulic press.",
                            "D) A battery charger.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 3,
                        "q": "What did the teacher tell them to inspect before replacing parts?",
                        "options": [
                            "A) The brake discs.",
                            "B) The cooling system.",
                            "C) The gearbox mount.",
                            "D) The exhaust pipe.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 4,
                        "q": "What was the real cause of the warning?",
                        "options": [
                            "A) A damaged battery pack.",
                            "B) A broken dashboard display.",
                            "C) A dusty filter that blocked the battery fan.",
                            "D) An empty coolant reservoir.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 5,
                        "q": "What conclusion did the group reach after the test drive?",
                        "options": [
                            "A) A simple blockage can create a serious warning signal.",
                            "B) Hybrid cars always need a new battery when a code appears.",
                            "C) The scan tool should not be used in class.",
                            "D) Cooling systems are less important than dashboard lights.",
                        ],
                        "answer": "A",
                    },
                ],
            },
            {
                "title": "PART II — Maintenance Procedure",
                "instructions": "Questions 6–10. Read the brake maintenance procedure and identify the correct technical sequence.",
                "text_title": "Text 2 — Brake Pad Replacement Procedure",
                "text": (
                    "The brake pad procedure begins after the vehicle is lifted safely and the wheel is removed. The technician loosens the "
                    "caliper bolts, compresses the piston with the correct tool, and removes the old pads. Before the new pads are installed, "
                    "the bracket is cleaned and the sliding pins are checked for smooth movement. After that, the new pads are placed in "
                    "position, the caliper is reassembled, and the bolts are tightened to the torque specified in the manual. Before the car "
                    "leaves the lift, the driver pumps the brake pedal several times to restore pressure."
                ),
                "questions": [
                    {
                        "num": 6,
                        "q": "What happens after the vehicle is lifted safely?",
                        "options": [
                            "A) The battery is disconnected.",
                            "B) The wheel is removed.",
                            "C) The engine oil is drained.",
                            "D) The tyres are balanced.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 7,
                        "q": "Why is the caliper piston compressed?",
                        "options": [
                            "A) To create space for the new pads.",
                            "B) To increase engine power.",
                            "C) To check the tyre pressure.",
                            "D) To clean the brake fluid reservoir.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 8,
                        "q": "What must be checked before the new pads are installed?",
                        "options": [
                            "A) The horn and the headlamps.",
                            "B) The bracket cleanliness and sliding pin movement.",
                            "C) The fuel filter and the spark plugs.",
                            "D) The windscreen washer level.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 9,
                        "q": "How should the caliper bolts be tightened?",
                        "options": [
                            "A) By hand only.",
                            "B) With the highest force possible.",
                            "C) According to the torque specified in the manual.",
                            "D) Only after the road test.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 10,
                        "q": "Why does the driver pump the brake pedal before leaving the lift?",
                        "options": [
                            "A) To restore brake pressure.",
                            "B) To warm up the engine.",
                            "C) To calibrate the dashboard screen.",
                            "D) To reduce wheel alignment.",
                        ],
                        "answer": "A",
                    },
                ],
            },
            {
                "title": "PART III — Safety Zone Notice",
                "instructions": "Questions 11–15. Read the safety notice and identify the correct precautions and the main risk being prevented in the EV work area.",
                "text_title": "Text 3 — Electric Vehicle Safety Zone",
                "text": (
                    "ELECTRIC VEHICLE AREA: Wear insulated gloves when checking orange high-voltage cables. Place warning cones around the car "
                    "before opening the service disconnect cover. Do not carry metal tools in your pocket while working near battery connectors. "
                    "If a cable shows damaged insulation, stop the inspection and inform the teacher immediately. Complete the final checklist "
                    "before reconnecting the system."
                ),
                "questions": [
                    {
                        "num": 11,
                        "q": "What should students wear when checking high-voltage cables?",
                        "options": [
                            "A) Welding boots only.",
                            "B) Insulated gloves.",
                            "C) Ear protection only.",
                            "D) Cloth gloves from the cleaning area.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 12,
                        "q": "What must be placed around the car before opening the cover?",
                        "options": [
                            "A) Spare tyres.",
                            "B) Warning cones.",
                            "C) Brake pads.",
                            "D) Cleaning cloths.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 13,
                        "q": "What should not be carried in a pocket near battery connectors?",
                        "options": [
                            "A) Plastic labels.",
                            "B) Metal tools.",
                            "C) A notebook.",
                            "D) Safety gloves.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 14,
                        "q": "What action is required if a cable has damaged insulation?",
                        "options": [
                            "A) Continue the inspection quickly.",
                            "B) Reconnect the system first.",
                            "C) Stop and inform the teacher immediately.",
                            "D) Hide the cable with tape and continue.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 15,
                        "q": "What overall risk does this notice mainly try to reduce?",
                        "options": [
                            "A) Electrical accidents during EV inspection.",
                            "B) Fuel consumption during road tests.",
                            "C) Noise from tyre balancing machines.",
                            "D) Delay in customer payments.",
                        ],
                        "answer": "A",
                    },
                ],
            },
        ],
    },
    {
        "file_stub": "PRUEBA_COMPRENSION_LECTORA_3RO_ELECTRICIDAD",
        "course": "3°C — Electricidad",
        "course_field": "3°C ____",
        "specialty": "Electricidad",
        "objective": "Evaluar la comprensión de textos técnicos vinculados a instalaciones, procedimientos de seguridad y control eléctrico.",
        "skills": "Comprensión lectora — vocabulario técnico — interpretación de diagramas y avisos",
        "parts": [
            {
                "title": "PART I — Installation Report",
                "instructions": "Questions 1–5. Read the report and choose the best answer. Each correct answer is worth 2 points.",
                "text_title": "Text 1 — LED Lighting Project",
                "text": (
                    "In the school workshop, an electricity group planned a new LED lighting circuit for a study room. The first survey showed "
                    "that the old fluorescent lamps used too much energy and produced uneven light. The students read the wiring diagram, "
                    "calculated the load, and selected two separate circuits: one for the main lights and another for emergency lighting. During "
                    "the installation, they labelled each conductor and checked the voltage with a multimeter before connecting the switches. "
                    "When the project was finished, the room was brighter and the panel was easier to understand during maintenance."
                ),
                "questions": [
                    {
                        "num": 1,
                        "q": "What was the purpose of the new project?",
                        "options": [
                            "A) To repair the school internet network.",
                            "B) To install a new LED lighting circuit in a study room.",
                            "C) To replace the laboratory computers.",
                            "D) To move the electrical panel to another building.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 2,
                        "q": "What problem did the old fluorescent lamps cause?",
                        "options": [
                            "A) They used too much energy and produced uneven light.",
                            "B) They disconnected the emergency exit signs.",
                            "C) They made the multimeter stop working.",
                            "D) They blocked the wiring diagram completely.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 3,
                        "q": "Why did the group select two separate circuits?",
                        "options": [
                            "A) To connect one circuit to the internet.",
                            "B) To separate main lighting from emergency lighting.",
                            "C) To avoid using any labels on conductors.",
                            "D) To increase the classroom temperature.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 4,
                        "q": "Which tool was used to check voltage before connecting the switches?",
                        "options": [
                            "A) A welding helmet.",
                            "B) A multimeter.",
                            "C) A paint roller.",
                            "D) A brake scanner.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 5,
                        "q": "What positive result is mentioned at the end of the text?",
                        "options": [
                            "A) The room became brighter and the panel was easier to understand.",
                            "B) The group removed the emergency lights completely.",
                            "C) The old lamps produced more heat than before.",
                            "D) The switches were connected without measuring voltage.",
                        ],
                        "answer": "A",
                    },
                ],
            },
            {
                "title": "PART II — Isolation Procedure",
                "instructions": "Questions 6–10. Read the procedure and identify the correct sequence for safe work.",
                "text_title": "Text 2 — Breaker Replacement Procedure",
                "text": (
                    "Before replacing a miniature breaker, the technician reviews the panel diagram and confirms which circuit must be isolated. "
                    "Then the main supply is switched off, a lockout device is placed on the handle, and a warning tag is added. After that, the "
                    "technician uses a tester to confirm that no voltage is present. Only then is the damaged breaker removed and the new unit "
                    "installed. At the end, the connections are tightened, the panel cover is replaced, and the circuit is tested again under supervision."
                ),
                "questions": [
                    {
                        "num": 6,
                        "q": "What is checked before the technician touches the breaker?",
                        "options": [
                            "A) The panel diagram and the circuit to isolate.",
                            "B) The school timetable.",
                            "C) The paint colour of the panel.",
                            "D) The classroom attendance list.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 7,
                        "q": "What is placed on the handle after the main supply is switched off?",
                        "options": [
                            "A) A cleaning cloth.",
                            "B) A lockout device.",
                            "C) A new conductor label.",
                            "D) A battery charger.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 8,
                        "q": "Why is a tester used before the damaged breaker is removed?",
                        "options": [
                            "A) To confirm that no voltage is present.",
                            "B) To measure room temperature.",
                            "C) To print a maintenance invoice.",
                            "D) To connect the emergency circuit.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 9,
                        "q": "What happens after the new breaker is installed?",
                        "options": [
                            "A) The wiring diagram is deleted.",
                            "B) The connections are tightened and the cover is replaced.",
                            "C) The voltage is left untested.",
                            "D) The old breaker is reinstalled.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 10,
                        "q": "How is the final circuit test performed?",
                        "options": [
                            "A) Without supervision.",
                            "B) Only after removing the cover forever.",
                            "C) Again under supervision.",
                            "D) Before the new breaker is installed.",
                        ],
                        "answer": "C",
                    },
                ],
            },
            {
                "title": "PART III — Safety Notice",
                "instructions": "Questions 11–15. Read the notice and identify the key precautions and the main risk being prevented in the work area.",
                "text_title": "Text 3 — Temporary Wiring Safety Notice",
                "text": (
                    "SAFETY NOTICE: Keep extension cords away from wet floors and doorways. Replace any conductor that shows broken insulation. "
                    "Do not overload a multi-plug adapter with several high-consumption devices. Before using a portable drill, inspect the plug "
                    "and confirm that the ground connection is intact. Report any unusual smell or excessive heat to the teacher immediately."
                ),
                "questions": [
                    {
                        "num": 11,
                        "q": "Where should extension cords be kept away from?",
                        "options": [
                            "A) Wet floors and doorways.",
                            "B) Clean notebooks only.",
                            "C) The lighting diagram only.",
                            "D) The teacher's desk only.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 12,
                        "q": "What should be replaced immediately?",
                        "options": [
                            "A) Any conductor with broken insulation.",
                            "B) Every new LED lamp.",
                            "C) The panel cover after each class.",
                            "D) Every lockout tag.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 13,
                        "q": "What should not be overloaded according to the notice?",
                        "options": [
                            "A) The emergency circuit.",
                            "B) A multi-plug adapter.",
                            "C) The study room window.",
                            "D) The grounding conductor.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 14,
                        "q": "What must be confirmed before using a portable drill?",
                        "options": [
                            "A) The battery level of the multimeter.",
                            "B) That the ground connection is intact.",
                            "C) The colour of the extension cord.",
                            "D) The age of the classroom panel.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 15,
                        "q": "What overall problem does this notice mainly help prevent?",
                        "options": [
                            "A) Electric shock or overheating from unsafe temporary wiring.",
                            "B) Incorrect student attendance.",
                            "C) Lack of tools in storage.",
                            "D) Poor internet signal in the workshop.",
                        ],
                        "answer": "A",
                    },
                ],
            },
        ],
    },
    {
        "file_stub": "PRUEBA_COMPRENSION_LECTORA_3RO_GRAFICA",
        "course": "3°D — Gráfica",
        "course_field": "3°D ____",
        "specialty": "Gráfica",
        "objective": "Evaluar la comprensión de textos técnicos vinculados a preprensa, impresión y protocolos de seguridad en Gráfica.",
        "skills": "Comprensión lectora — vocabulario técnico — interpretación de procesos de impresión",
        "parts": [
            {
                "title": "PART I — Prepress Case",
                "instructions": "Questions 1–5. Read the case and choose the best answer. Each correct answer is worth 2 points.",
                "text_title": "Text 1 — Colour Mismatch in a Poster Run",
                "text": (
                    "A graphic design group prepared a poster for a school campaign, but the printed result looked darker than the digital proof "
                    "on screen. To solve the problem, the students compared the CMYK profile, the paper type, and the printer settings. They "
                    "discovered that the file had been exported with the wrong colour profile and that the monitor had not been calibrated recently. "
                    "After correcting the profile and printing a new proof, the images looked closer to the expected colours. The teacher used the "
                    "case to explain that professional printing depends on careful prepress decisions, not only on creative ideas."
                ),
                "questions": [
                    {
                        "num": 1,
                        "q": "What problem appeared in the first print?",
                        "options": [
                            "A) The poster was brighter than expected.",
                            "B) The poster looked darker than the digital proof.",
                            "C) The printer stopped because of a broken wheel.",
                            "D) The paper was too small for the design.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 2,
                        "q": "Which technical elements did the students compare?",
                        "options": [
                            "A) The CMYK profile, paper type, and printer settings.",
                            "B) The classroom schedule and student list.",
                            "C) The website password and USB cable.",
                            "D) The wall colour and poster frame.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 3,
                        "q": "What error was found in the exported file?",
                        "options": [
                            "A) It had no title on the first page.",
                            "B) It used the wrong colour profile.",
                            "C) It was saved as an audio file.",
                            "D) It contained no images at all.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 4,
                        "q": "What happened after the profile was corrected?",
                        "options": [
                            "A) The new proof looked closer to the expected colours.",
                            "B) The printer produced only black pages.",
                            "C) The campaign had to be cancelled.",
                            "D) The monitor stopped showing any image.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 5,
                        "q": "What professional lesson does the teacher highlight?",
                        "options": [
                            "A) Creative ideas are the only important part of printing.",
                            "B) Professional printing depends on careful prepress decisions.",
                            "C) Paper type never affects colour results.",
                            "D) Colour profiles are not used in graphic design.",
                        ],
                        "answer": "B",
                    },
                ],
            },
            {
                "title": "PART II — Export Procedure",
                "instructions": "Questions 6–10. Read the procedure and identify the correct prepress sequence.",
                "text_title": "Text 2 — PDF for Print Export Checklist",
                "text": (
                    "Before sending a poster to print, the operator runs a preflight check on the file. First, the bleed is confirmed on all "
                    "sides and crop marks are activated. Next, the images are reviewed to make sure they have enough resolution for printing. "
                    "The fonts are embedded, and the colours are converted to CMYK when necessary. After that, the operator exports the PDF and "
                    "prints a small proof to check margins, colour balance, and text readability. Only after that proof is approved does the file "
                    "go to final production."
                ),
                "questions": [
                    {
                        "num": 6,
                        "q": "What is checked first in the export process?",
                        "options": [
                            "A) The classroom attendance.",
                            "B) The bleed and crop marks.",
                            "C) The air temperature in the room.",
                            "D) The printer wheels.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 7,
                        "q": "Why are the images reviewed before export?",
                        "options": [
                            "A) To confirm that they have enough resolution for printing.",
                            "B) To turn them into handwritten sketches.",
                            "C) To remove every colour from the design.",
                            "D) To convert them into spreadsheets.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 8,
                        "q": "What happens to the fonts in the file?",
                        "options": [
                            "A) They are erased before export.",
                            "B) They are embedded.",
                            "C) They are replaced with random symbols.",
                            "D) They are printed by hand.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 9,
                        "q": "Why is a small proof printed after export?",
                        "options": [
                            "A) To check margins, colour balance, and readability.",
                            "B) To make the final production slower.",
                            "C) To delete the crop marks automatically.",
                            "D) To avoid using CMYK colours.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 10,
                        "q": "When is the file sent to final production?",
                        "options": [
                            "A) Before any proof is printed.",
                            "B) Only after the proof is approved.",
                            "C) As soon as the bleed is ignored.",
                            "D) After the fonts are removed.",
                        ],
                        "answer": "B",
                    },
                ],
            },
            {
                "title": "PART III — Print Room Notice",
                "instructions": "Questions 11–15. Read the notice and identify the correct safety and cleaning instructions, including its main purpose.",
                "text_title": "Text 3 — Roller Cleaning Notice",
                "text": (
                    "PRINT ROOM NOTICE: Use gloves and ventilation when cleaning ink rollers with solvent. Keep open containers away from heat "
                    "sources and close them immediately after use. Store wiping cloths in the metal container marked for chemical waste. If a "
                    "spill reaches the floor, block the area and report it before starting another task. Wash your hands before touching paper, "
                    "proofs, or the keyboard."
                ),
                "questions": [
                    {
                        "num": 11,
                        "q": "What should be used when cleaning rollers with solvent?",
                        "options": [
                            "A) Gloves and ventilation.",
                            "B) Water and sand only.",
                            "C) Bare hands and a fan turned off.",
                            "D) Brake fluid and paper towels.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 12,
                        "q": "Where must open solvent containers be kept away from?",
                        "options": [
                            "A) The paper cutter only.",
                            "B) Heat sources.",
                            "C) The keyboard only.",
                            "D) Crop marks only.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 13,
                        "q": "Where should wiping cloths be stored after use?",
                        "options": [
                            "A) Inside a student backpack.",
                            "B) In the metal container for chemical waste.",
                            "C) Under the printer.",
                            "D) On the nearest table.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 14,
                        "q": "What should happen if a spill reaches the floor?",
                        "options": [
                            "A) Start another task immediately.",
                            "B) Ignore it if the floor is already dark.",
                            "C) Block the area and report the spill.",
                            "D) Print a proof over the spill.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 15,
                        "q": "What is the main purpose of this notice in the print room?",
                        "options": [
                            "A) To prevent chemical accidents and contamination during cleaning.",
                            "B) To speed up poster export and proof approval.",
                            "C) To reduce the number of keyboards in the room.",
                            "D) To replace all solvent with water.",
                        ],
                        "answer": "A",
                    },
                ],
            },
        ],
    },
    {
        "file_stub": "PRUEBA_COMPRENSION_LECTORA_3RO_ELECTRONICA",
        "course": "3°E — Electrónica",
        "course_field": "3°E ____",
        "specialty": "Electrónica",
        "objective": "Evaluar la comprensión de textos técnicos vinculados a circuitos, procedimientos de soldadura y seguridad electrónica.",
        "skills": "Comprensión lectora — vocabulario técnico — interpretación de fichas y protocolos",
        "parts": [
            {
                "title": "PART I — Circuit Troubleshooting",
                "instructions": "Questions 1–5. Read the report and choose the best answer. Each correct answer is worth 2 points.",
                "text_title": "Text 1 — Debugging a Microcontroller Circuit",
                "text": (
                    "In the electronics lab, a team built a small circuit to make an LED blink with a microcontroller. When they uploaded the code, "
                    "the LED stayed off. First, they checked the program, but the commands were correct. Then they used a multimeter and compared "
                    "the circuit with the datasheet. They noticed that the resistor value was too high and that the ground cable was loose on the "
                    "breadboard. After replacing the resistor and fixing the connection, the LED started blinking normally. The group learned that "
                    "debugging requires both software review and careful inspection of the physical circuit."
                ),
                "questions": [
                    {
                        "num": 1,
                        "q": "What was the original problem with the circuit?",
                        "options": [
                            "A) The LED stayed off after the code was uploaded.",
                            "B) The multimeter stopped working.",
                            "C) The breadboard was missing completely.",
                            "D) The keyboard could not type the program.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 2,
                        "q": "What did the students check first?",
                        "options": [
                            "A) The soldering station temperature.",
                            "B) The program commands.",
                            "C) The classroom lights.",
                            "D) The air conditioner.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 3,
                        "q": "Which tools or documents helped them find the problem?",
                        "options": [
                            "A) A multimeter and the datasheet.",
                            "B) A printer and a paint brush.",
                            "C) A tyre gauge and a jack.",
                            "D) A welding helmet and a ladder.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 4,
                        "q": "What two technical faults were discovered?",
                        "options": [
                            "A) A low battery and a broken monitor.",
                            "B) A resistor value that was too high and a loose ground cable.",
                            "C) A missing teacher and a noisy fan.",
                            "D) An empty solder spool and a cracked table.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 5,
                        "q": "What lesson did the group learn at the end?",
                        "options": [
                            "A) Debugging only requires rewriting the code.",
                            "B) Physical circuits never need inspection.",
                            "C) Debugging needs software review and careful circuit inspection.",
                            "D) LEDs should not be used with microcontrollers.",
                        ],
                        "answer": "C",
                    },
                ],
            },
            {
                "title": "PART II — Workshop Procedure",
                "instructions": "Questions 6–10. Read the procedure and identify the correct soldering sequence.",
                "text_title": "Text 2 — Soldering a Header to a PCB",
                "text": (
                    "Before soldering a header to a printed circuit board, the technician cleans the tip of the soldering iron and fixes the "
                    "board in a small support. The pins are inserted into the correct holes, and the component is held in place. Then the pad "
                    "and the pin are heated at the same time for a short moment. After that, solder is applied until a smooth, shiny joint appears. "
                    "When the line of pins is finished, the technician inspects each joint and removes any solder bridge before powering the circuit."
                ),
                "questions": [
                    {
                        "num": 6,
                        "q": "What is cleaned before the soldering process starts?",
                        "options": [
                            "A) The monitor screen.",
                            "B) The soldering iron tip.",
                            "C) The power cable insulation.",
                            "D) The keyboard battery.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 7,
                        "q": "Why is the board fixed in a support?",
                        "options": [
                            "A) To stop the classroom lights from moving.",
                            "B) To hold the board steady during the task.",
                            "C) To cool the solder faster.",
                            "D) To replace the power supply.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 8,
                        "q": "What are heated at the same time before solder is added?",
                        "options": [
                            "A) The pad and the pin.",
                            "B) The datasheet and the keyboard.",
                            "C) The PCB box and the workbench.",
                            "D) The switch and the multimeter screen.",
                        ],
                        "answer": "A",
                    },
                    {
                        "num": 9,
                        "q": "How should a correct joint look according to the text?",
                        "options": [
                            "A) Dark and rough.",
                            "B) Large and melted across two pads.",
                            "C) Smooth and shiny.",
                            "D) Hidden under the connector.",
                        ],
                        "answer": "C",
                    },
                    {
                        "num": 10,
                        "q": "What must be removed before the circuit is powered?",
                        "options": [
                            "A) The component support.",
                            "B) Any solder bridge.",
                            "C) The board labels.",
                            "D) The ground connection.",
                        ],
                        "answer": "B",
                    },
                ],
            },
            {
                "title": "PART III — Lab Safety Notice",
                "instructions": "Questions 11–15. Read the notice and identify the correct ESD precautions and the risk they prevent.",
                "text_title": "Text 3 — Electrostatic Discharge Notice",
                "text": (
                    "ESD NOTICE: Use the anti-static wrist strap whenever handling integrated circuits outside their protective bags. Keep component "
                    "boxes closed when they are not in use. Do not place microchips directly on metal tables or plastic folders. If the wrist strap "
                    "fails the tester, replace it before continuing the task. Return sensitive components to their labelled containers after each activity."
                ),
                "questions": [
                    {
                        "num": 11,
                        "q": "When must the anti-static wrist strap be used?",
                        "options": [
                            "A) Only while cleaning the floor.",
                            "B) Whenever integrated circuits are handled outside their bags.",
                            "C) Only after the class has finished.",
                            "D) Only for mechanical repairs.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 12,
                        "q": "What should happen to component boxes when they are not being used?",
                        "options": [
                            "A) They should stay open all day.",
                            "B) They should be closed.",
                            "C) They should be placed on the floor.",
                            "D) They should be washed with solvent.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 13,
                        "q": "Where should microchips not be placed?",
                        "options": [
                            "A) On a labelled foam tray.",
                            "B) On metal tables or plastic folders.",
                            "C) In protective bags.",
                            "D) In closed component boxes.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 14,
                        "q": "What must happen if the wrist strap fails the tester?",
                        "options": [
                            "A) Continue the task quickly.",
                            "B) Replace the strap before continuing.",
                            "C) Remove all labels from the boxes.",
                            "D) Power the circuit immediately.",
                        ],
                        "answer": "B",
                    },
                    {
                        "num": 15,
                        "q": "What overall risk does this notice mainly try to reduce?",
                        "options": [
                            "A) Damage to components from electrostatic discharge and poor handling.",
                            "B) Incorrect tyre pressure in the workshop.",
                            "C) Overheating of the soldering iron tip.",
                            "D) Colour mismatch in printed posters.",
                        ],
                        "answer": "A",
                    },
                ],
            },
        ],
    },
]


def set_cell_shading(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    shading = cell_properties.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    cell_properties.append(shading)


def add_part_header(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    paragraph_properties = paragraph._element.get_or_add_pPr()
    shading = paragraph_properties.makeelement(qn("w:shd"), {qn("w:fill"): "2962FF", qn("w:val"): "clear"})
    paragraph_properties.append(shading)


def add_instruction_line(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_text_block(document, title, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph_properties = paragraph._element.get_or_add_pPr()
    shading = paragraph_properties.makeelement(qn("w:shd"), {qn("w:fill"): "F5F5F5", qn("w:val"): "clear"})
    paragraph_properties.append(shading)
    title_run = paragraph.add_run(title + "\n")
    title_run.bold = True
    title_run.font.size = Pt(9)
    text_run = paragraph.add_run(text)
    text_run.font.size = Pt(8.5)


def add_question(document, number, question_text, options):
    question = document.add_paragraph()
    question.paragraph_format.space_before = Pt(2)
    question.paragraph_format.space_after = Pt(1)
    question_run = question.add_run(f"{number}. {question_text}")
    question_run.bold = True
    question_run.font.size = Pt(8.5)
    for option in options:
        option_paragraph = document.add_paragraph()
        option_paragraph.paragraph_format.space_before = Pt(0)
        option_paragraph.paragraph_format.space_after = Pt(0)
        option_paragraph.paragraph_format.left_indent = Cm(0.5)
        option_run = option_paragraph.add_run(option)
        option_run.font.size = Pt(8)


def build_score_rows(total_points):
    rows = []
    for points in range(0, total_points + 1):
        if points < PASSING_POINTS:
            grade = 1.0 + 3.0 * (points / PASSING_POINTS)
        else:
            grade = 4.0 + 3.0 * ((points - PASSING_POINTS) / (total_points - PASSING_POINTS))
        rows.append((points, round(min(grade, 7.0), 1)))
    return rows


def all_answers(assessment):
    answers = {}
    for part in assessment["parts"]:
        for question in part["questions"]:
            answers[question["num"]] = question["answer"]
    return answers


def generate_html(assessment):
    answers = all_answers(assessment)
    parts_html = []
    for part in assessment["parts"]:
        question_html = []
        for question in part["questions"]:
            options_html = "".join(f"<div class=\"option\">{option}</div>" for option in question["options"])
            question_html.append(
                f"<div class=\"question\"><p class=\"q-text\">{question['num']}. {question['q']}</p>{options_html}</div>"
            )
        parts_html.append(
            "\n".join(
                [
                    '<div class="part">',
                    f'<div class="part-header">{part["title"]} (10 points)</div>',
                    f'<p class="instructions">{part["instructions"]}</p>',
                    f'<div class="reading-text"><span class="text-title">{part["text_title"]}</span><br/>{part["text"]}</div>',
                    *question_html,
                    "</div>",
                ]
            )
        )

    answer_items = "".join(
        f'<div class="answer-item"><span class="ans-num">{number}.</span> {answers[number]}</div>'
        for number in sorted(answers)
    )

    score_rows = []
    scores = build_score_rows(TOTAL_POINTS)
    rows_per_column = (len(scores) + 2) // 3
    for row_index in range(rows_per_column):
        row_cells = []
        for column in range(3):
            score_index = row_index + column * rows_per_column
            if score_index < len(scores):
                points, grade = scores[score_index]
                highlight = ' style="background:#e8f5e9;"' if grade >= 4.0 else ""
                row_cells.append(f"<td>{points}</td><td{highlight}>{grade}</td>")
            else:
                row_cells.append("<td></td><td></td>")
        score_rows.append("<tr>" + "".join(row_cells) + "</tr>")

    logo_html = f'<img src="data:image/png;base64,{LOGO_B64}" alt="Logo Liceo"/>' if LOGO_B64 else ""

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8"/>
<meta content="width=device-width, initial-scale=1.0" name="viewport"/>
<title>{assessment['course']} - Prueba de Comprension Lectora TP</title>
<style>
    @page {{
        size: 21.59cm 27.94cm;
        margin: 1.5cm 1.8cm 1.5cm 1.8cm;
    }}
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{
        font-family: Arial, Helvetica, sans-serif;
        font-size: 9pt;
        line-height: 1.3;
        color: #1a1a1a;
        background: #f0f0f0;
    }}
    .page {{
        background: white;
        max-width: 21.59cm;
        margin: 1cm auto;
        padding: 0.5cm 1.8cm 1.5cm 1.8cm;
        box-shadow: 0 2px 12px rgba(0,0,0,0.12);
    }}
    .logo-container img {{ width: 100%; display: block; margin-bottom: -4px; }}
    .test-title {{ text-align: center; font-size: 11pt; font-weight: bold; margin: 4px 0; }}
    .field-table, .criteria-box, .score-table {{ width: 100%; border-collapse: collapse; }}
    .field-table td, .criteria-box td, .score-table td, .score-table th {{ border: 1px solid #000; padding: 2px 5px; }}
    .label {{ background: #e8e8e8; font-weight: bold; }}
    .criteria-box {{ margin-bottom: 6px; font-size: 8pt; }}
    .part {{ margin-bottom: 8px; }}
    .part-header {{ background: #2962FF; color: white; padding: 3px 8px; font-weight: bold; font-size: 9pt; margin-bottom: 3px; }}
    .instructions {{ font-style: italic; font-size: 8pt; color: #333; margin-bottom: 4px; line-height: 1.35; }}
    .reading-text {{ background: #f5f5f5; border: 1px solid #ddd; padding: 6px 8px; margin-bottom: 6px; font-size: 8.5pt; line-height: 1.35; }}
    .text-title {{ font-weight: bold; font-size: 9pt; }}
    .question {{ margin-bottom: 4px; }}
    .q-text {{ font-weight: bold; font-size: 8.5pt; margin-bottom: 1px; }}
    .option {{ padding-left: 14px; font-size: 8pt; line-height: 1.35; }}
    .test-footer {{ text-align: center; font-style: italic; font-size: 7.5pt; color: #888; margin-top: 10px; padding-top: 6px; border-top: 1px solid #ddd; }}
    .answer-key {{ page-break-before: always; }}
    .answer-key-title {{ text-align: center; font-size: 10pt; font-weight: bold; margin-bottom: 6px; }}
    .answer-key-subtitle {{ text-align: center; font-size: 8.5pt; color: #555; margin-bottom: 8px; }}
    .answer-grid {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 2px 12px; font-size: 8pt; }}
    .answer-item .ans-num {{ font-weight: bold; display: inline-block; min-width: 22px; }}
    .score-table {{ margin-top: 18px; font-size: 8pt; }}
    .score-table th {{ background: #e8e8e8; }}
    @media print {{
        body {{ background: white; }}
        .page {{ max-width: none; margin: 0; padding: 0; box-shadow: none; }}
        .part-header, .label {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    }}
</style>
</head>
<body>
<div class="page">
    <div class="logo-container">{logo_html}</div>
    <div class="test-title">Prueba de Comprension Lectora - Ingles TP - {assessment['course']}</div>
    <table class="field-table">
        <tr><td class="label">Profesor:</td><td></td><td class="label">Nombre</td><td colspan="3"></td></tr>
        <tr><td class="label">Asignatura:</td><td>Idioma Extranjero: Ingles</td><td class="label">Curso</td><td colspan="3">{assessment['course_field']}</td></tr>
        <tr><td class="label">Prueba N°</td><td>Nota 2 - Semestre 1</td><td class="label">N° Lista</td><td colspan="3"></td></tr>
        <tr><td class="label">Semestre</td><td>1</td><td class="label">Puntos Obtenidos</td><td></td><td class="label">NOTA</td><td></td></tr>
        <tr><td class="label">Fecha</td><td colspan="5">___/___/2026</td></tr>
        <tr><td class="label">Visado por:</td><td colspan="5">EPM Gabriel S. Castro A</td></tr>
        <tr><td class="label">Objetivo:</td><td colspan="5">{assessment['objective']}</td></tr>
        <tr><td class="label">Habilidades</td><td colspan="5">{assessment['skills']}</td></tr>
    </table>
    <table class="criteria-box">
        <tr>
            <td class="label">Instrucciones:</td>
            <td>
                • Lee con atencion cada texto antes de responder.<br/>
                • Cada respuesta correcta vale 2 puntos.<br/>
                • Marca la alternativa correcta entre A, B, C o D.<br/>
                • Puntaje total: {TOTAL_POINTS} puntos. Exigencia: 60%.<br/>
                • Tiempo sugerido: 60 minutos.<br/>
                • No se permite el uso de diccionario ni celular.
            </td>
        </tr>
    </table>
    {''.join(parts_html)}
    <div class="test-footer">- Fin de la evaluacion - Revisa tus respuestas antes de entregar. -</div>
</div>

<div class="page answer-key">
    <div class="answer-key-title">Pauta de Respuestas Correctas</div>
    <div class="answer-key-subtitle">Prueba de Comprension Lectora - {assessment['course']}</div>
    <div class="answer-grid">{answer_items}</div>
    <table class="score-table">
        <tr><th colspan="6">Tabla de Conversion de Puntaje a Nota (60% de exigencia)</th></tr>
        <tr><th>Puntaje</th><th>Nota</th><th>Puntaje</th><th>Nota</th><th>Puntaje</th><th>Nota</th></tr>
        {''.join(score_rows)}
    </table>
</div>
</body>
</html>"""


def generate_docx(assessment, output_path):
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Cm(17.99))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(f"Prueba de Comprension Lectora - Ingles TP - {assessment['course']}")
    title_run.bold = True
    title_run.font.size = Pt(11)

    table = document.add_table(rows=8, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    def set_cell(row, col, text, is_label=False, colspan=1):
        cell = table.cell(row, col)
        if colspan > 1:
            cell.merge(table.cell(row, col + colspan - 1))
        cell.text = text
        if is_label:
            set_cell_shading(cell, "E8E8E8")
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8)
                if is_label:
                    run.bold = True

    set_cell(0, 0, "Profesor:", True)
    set_cell(0, 1, "")
    set_cell(0, 2, "Nombre", True)
    set_cell(0, 3, "", colspan=3)
    set_cell(1, 0, "Asignatura:", True)
    set_cell(1, 1, "Idioma Extranjero: Ingles")
    set_cell(1, 2, "Curso", True)
    set_cell(1, 3, assessment["course_field"], colspan=3)
    set_cell(2, 0, "Prueba N°", True)
    set_cell(2, 1, "Nota 2 - Semestre 1")
    set_cell(2, 2, "N° Lista", True)
    set_cell(2, 3, "", colspan=3)
    set_cell(3, 0, "Semestre", True)
    set_cell(3, 1, "1")
    set_cell(3, 2, "Puntos Obtenidos", True)
    set_cell(3, 3, "")
    set_cell(3, 4, "NOTA", True)
    set_cell(3, 5, "")
    set_cell(4, 0, "Fecha", True)
    set_cell(4, 1, "___/___/2026", colspan=5)
    set_cell(5, 0, "Visado por:", True)
    set_cell(5, 1, "EPM Gabriel S. Castro A", colspan=5)
    set_cell(6, 0, "Objetivo:", True)
    set_cell(6, 1, assessment["objective"], colspan=5)
    set_cell(7, 0, "Habilidades", True)
    set_cell(7, 1, assessment["skills"], colspan=5)

    instructions = document.add_paragraph()
    instructions.paragraph_format.space_before = Pt(4)
    instructions.paragraph_format.space_after = Pt(4)
    instructions_run = instructions.add_run("Instrucciones:\n")
    instructions_run.bold = True
    instructions_run.font.size = Pt(8)
    instructions_text = (
        "• Lee con atencion cada texto antes de responder.\n"
        "• Cada respuesta correcta vale 2 puntos.\n"
        "• Marca la alternativa correcta entre A, B, C o D.\n"
        "• La prueba tiene 30 puntos en total. Nivel de exigencia: 60%.\n"
        "• Tienes 60 minutos para completar la prueba.\n"
        "• No se permite el uso de diccionario ni dispositivos electronicos."
    )
    instructions.add_run(instructions_text).font.size = Pt(7.5)

    for part in assessment["parts"]:
        add_part_header(document, part["title"] + " (10 points)")
        add_instruction_line(document, part["instructions"])
        add_text_block(document, part["text_title"], part["text"])
        for question in part["questions"]:
            add_question(document, question["num"], question["q"], question["options"])

    document.add_page_break()

    answer_title = document.add_paragraph()
    answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_title_run = answer_title.add_run("Pauta de Respuestas Correctas")
    answer_title_run.bold = True
    answer_title_run.font.size = Pt(11)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Prueba de Comprension Lectora - {assessment['course']}")
    subtitle_run.font.size = Pt(8.5)
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    answer_table = document.add_table(rows=0, cols=5)
    answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    answers = sorted(all_answers(assessment).items())
    for index in range(0, len(answers), 5):
        row = answer_table.add_row()
        for column, item in enumerate(answers[index:index + 5]):
            cell = row.cells[column]
            cell.text = f"{item[0]}. {item[1]}"
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.bold = True

    score_table = document.add_table(rows=2, cols=6)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_table.style = "Table Grid"
    score_table.cell(0, 0).merge(score_table.cell(0, 5))
    score_table.cell(0, 0).text = "Tabla de Conversion de Puntaje a Nota (60% de exigencia)"
    for paragraph in score_table.cell(0, 0).paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)
    headers = ["Puntaje", "Nota", "Puntaje", "Nota", "Puntaje", "Nota"]
    for column, header in enumerate(headers):
        score_table.cell(1, column).text = header
        set_cell_shading(score_table.cell(1, column), "E8E8E8")
        for paragraph in score_table.cell(1, column).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)

    scores = build_score_rows(TOTAL_POINTS)
    rows_per_column = (len(scores) + 2) // 3
    for row_index in range(rows_per_column):
        row = score_table.add_row()
        for column in range(3):
            score_index = row_index + column * rows_per_column
            if score_index < len(scores):
                points, grade = scores[score_index]
                points_cell = row.cells[column * 2]
                grade_cell = row.cells[column * 2 + 1]
                points_cell.text = str(points)
                grade_cell.text = str(grade)
                if grade >= 4.0:
                    set_cell_shading(grade_cell, "E8F5E9")
                for cell in (points_cell, grade_cell):
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(7)

    document.save(output_path)


def write_targets(path, content):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def publish_generated_files(file_stub, html_content, docx_path):
    html_plan_path = PLAN_OUTPUT_DIR / f"{file_stub}.html"
    docx_plan_path = PLAN_OUTPUT_DIR / f"{file_stub}.docx"
    write_targets(html_plan_path, html_content)
    if docx_path != docx_plan_path:
        docx_plan_path.parent.mkdir(parents=True, exist_ok=True)
        docx_plan_path.write_bytes(docx_path.read_bytes())

    for target_dir in TARGET_DIRS:
        target_dir.mkdir(parents=True, exist_ok=True)
        (target_dir / f"{file_stub}.html").write_text(html_content, encoding="utf-8")
        (target_dir / f"{file_stub}.docx").write_bytes(docx_path.read_bytes())


def generate_all():
    PLAN_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for assessment in COURSE_ASSESSMENTS:
        html_content = generate_html(assessment)
        docx_output_path = PLAN_OUTPUT_DIR / f"{assessment['file_stub']}.docx"
        generate_docx(assessment, docx_output_path)
        publish_generated_files(assessment["file_stub"], html_content, docx_output_path)
        print(f"Generado: {assessment['file_stub']}")


if __name__ == "__main__":
    generate_all()