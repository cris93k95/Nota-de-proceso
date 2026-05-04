#!/usr/bin/env python3
"""
Genera una evaluacion especial institucional para 1ro Medio.

La evaluacion esta adaptada para un estudiante que no seguira asistiendo
presencialmente durante el primer semestre, por lo que incluye orientaciones
autonomas, trabajo de expresion escrita, trabajo de expresion oral y sus
respectivas pautas de evaluacion.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
EVALUATION_OUTPUT_NAME = "EVALUACION_ESPECIAL_SEMESTRE_1_1RO_MEDIO_EXPRESION_ESCRITA_ORAL.docx"
GUIDE_OUTPUT_NAME = "GUIA_DE_APOYO_EVALUACION_ESPECIAL_1RO_MEDIO.docx"
LOGO_PATH = BASE_DIR / "Logo Colegio.png"

COLORS = {
    "blue": "1A237E",
    "light_blue": "E8EAF6",
    "gray": "E8E8E8",
    "light_gray": "F5F5F5",
    "rose": "FCE4EC",
    "amber": "FFF3E0",
    "teal": "E3F2FD",
}


def build_output_paths(file_name):
    return [
        BASE_DIR / "PLANIFICACIONES_2026_LISTO_IMPRESION" / file_name,
        BASE_DIR / "tranquiprofe.cl" / "static" / "recursos" / "materiales" / "1ro-medio" / "instrumentos" / file_name,
    ]


EVALUATION_OUTPUT_PATHS = build_output_paths(EVALUATION_OUTPUT_NAME)
GUIDE_OUTPUT_PATHS = build_output_paths(GUIDE_OUTPUT_NAME)


def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    tc_pr.append(shd)


def style_paragraph(paragraph, *, size=9, bold=False, italic=False, color=None, align=None, space_before=0, space_after=0):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_logo(doc):
    if not LOGO_PATH.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(LOGO_PATH), width=Cm(17.8))
    paragraph.paragraph_format.space_after = Pt(4)


def add_title(doc, title, subtitle):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(COLORS["blue"])

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(9)


def set_table_text(cell, text, *, bold=False, size=8.5, color=None):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def add_info_table(
    doc,
    *,
    evaluation_name,
    due_date,
    total_points,
    products,
    objective,
    modality="Trabajo autonomo no presencial",
    exigency="60%",
):
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("Profesor:", "", "Nombre estudiante:", ""),
        ("Asignatura:", "Idioma Extranjero: Ingles", "Curso:", "1ro Medio"),
        ("Documento:", evaluation_name, "Fecha de entrega:", due_date),
        ("Modalidad:", modality, "Puntaje total:", total_points),
        ("Exigencia:", exigency, "Productos:", products),
        (
            "Objetivo:",
            objective,
            "",
            "",
        ),
    ]

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            set_table_text(table.cell(row_idx, col_idx), value, bold=(col_idx % 2 == 0 and row_idx < 5))
            if col_idx % 2 == 0 or (row_idx == 5 and col_idx == 0):
                set_cell_shading(table.cell(row_idx, col_idx), COLORS["gray"])

    table.cell(5, 1).merge(table.cell(5, 3))
    set_table_text(
        table.cell(5, 1),
        objective,
        size=8,
    )


def add_box_heading(doc, text, fill_color):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_pr = paragraph._element.get_or_add_pPr()
    shd = p_pr.makeelement(qn("w:shd"), {qn("w:fill"): fill_color, qn("w:val"): "clear"})
    p_pr.append(shd)


def add_body_paragraph(doc, text, *, size=8.5, italic=False, bold=False, left_indent=0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    if left_indent:
        paragraph.paragraph_format.left_indent = Cm(left_indent)
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return paragraph


def add_bullets(doc, items, *, left_indent=0.35):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(left_indent)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(item)
        run.font.size = Pt(8.5)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(item)
        run.font.size = Pt(8.5)


def add_support_table(doc, title, rows, fill):
    add_box_heading(doc, title, fill)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Recurso", "Apoyo para el estudiante"]
    for idx, header in enumerate(headers):
        set_table_text(table.cell(0, idx), header, bold=True, color="FFFFFF")
        set_cell_shading(table.cell(0, idx), COLORS["blue"])
    for left, right in rows:
        cells = table.add_row().cells
        set_table_text(cells[0], left, bold=True, size=8)
        set_table_text(cells[1], right, size=8)
        set_cell_shading(cells[0], fill)


def add_rubric_table(doc, title, criteria_rows):
    add_box_heading(doc, title, COLORS["blue"])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Criterio", "4 puntos", "3 puntos", "2 puntos", "1 punto"]
    for idx, header in enumerate(headers):
        set_table_text(table.cell(0, idx), header, bold=True, color="FFFFFF")
        set_cell_shading(table.cell(0, idx), COLORS["blue"])

    for row in criteria_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_table_text(cells[idx], value, bold=(idx == 0), size=8)
            if idx == 0:
                set_cell_shading(cells[idx], COLORS["light_blue"])


def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    return doc


def build_document():
    doc = create_document()

    add_logo(doc)
    add_title(
        doc,
        "Evaluacion Especial Semestre 1 - Ingles - 1ro Medio",
        "Expresion escrita y expresion oral en contexto tecnico-profesional",
    )
    add_info_table(
        doc,
        evaluation_name="Evaluacion Especial Semestre 1",
        due_date="____ / ____ / 2026",
        total_points="40 puntos",
        products="1 trabajo escrito + 1 trabajo oral",
        objective="Evaluar expresion escrita y expresion oral integrando vocabulario tecnico TP, Present Perfect, since/for y used to, en un formato accesible para realizacion autonoma.",
    )

    add_box_heading(doc, "1. Presentacion e instrucciones generales", COLORS["blue"])
    add_body_paragraph(
        doc,
        "Este instrumento esta pensado para un estudiante que continuara el semestre en modalidad no presencial. Por ello, ambas tareas incluyen apoyos, ejemplos y una pauta clara para que puedan resolverse de forma autonoma, manteniendo el foco en los aprendizajes del semestre.",
    )
    add_bullets(
        doc,
        [
            "La evaluacion considera dos productos: un trabajo de expresion escrita y un trabajo de expresion oral.",
            "Ambos trabajos deben realizarse en ingles, pero puedes planificar tus ideas en espanol antes de escribir o grabar.",
            "Se permite usar diccionario, cuaderno y este mismo documento como apoyo. El producto final debe ser personal y original.",
            "El trabajo oral puede enviarse como audio o video, segun la indicacion del docente.",
            "Si tienes dificultades de conectividad, el trabajo escrito puede entregarse manuscrito con letra clara y el oral puede grabarse desde un telefono celular.",
        ],
    )

    add_support_table(
        doc,
        "2. Contenidos del semestre considerados en esta evaluacion",
        [
            (
                "Vocabulario tecnico TP",
                "Palabras y expresiones vinculadas a las cinco especialidades del colegio: Mecanica Automotriz, Mecanica Industrial, Electricidad, Electronica y Grafica.",
            ),
            (
                "Present Perfect",
                "Uso de have/has + past participle para hablar de experiencias, aprendizajes o procesos vinculados al mundo tecnico.",
            ),
            (
                "Since / For",
                "Uso para expresar duracion: since para un punto de inicio y for para un periodo de tiempo.",
            ),
            (
                "Used to",
                "Uso para comparar practicas del pasado con practicas actuales dentro de una especialidad tecnica.",
            ),
            (
                "Proyeccion personal",
                "Descripcion de la especialidad de interes, tareas, herramientas, importancia del ingles y futuro formativo o laboral.",
            ),
        ],
        COLORS["amber"],
    )

    add_box_heading(doc, "3. Trabajo 1 - Expresion escrita (20 puntos)", COLORS["blue"])
    add_body_paragraph(doc, "Titulo sugerido: My Future Technical Career", bold=True)
    add_body_paragraph(
        doc,
        "Escribe un texto en ingles de entre 180 y 220 palabras sobre la especialidad tecnica que mas te interesa o sobre la especialidad que te gustaria estudiar en el futuro.",
    )
    add_body_paragraph(doc, "Tu texto debe incluir obligatoriamente:", bold=True)
    add_numbered(
        doc,
        [
            "Una presentacion breve de ti mismo y de la especialidad elegida.",
            "Al menos 5 palabras de vocabulario tecnico relacionadas con una o mas especialidades TP del colegio.",
            "Al menos 2 oraciones en Present Perfect: una con since y una con for.",
            "Al menos 2 oraciones con used to para comparar como era el trabajo tecnico antes y como es ahora.",
            "Una conclusion donde expliques por que el ingles es importante para esa especialidad o para tu futuro laboral.",
        ],
    )
    add_body_paragraph(doc, "Sugerencia de estructura:", bold=True)
    add_bullets(
        doc,
        [
            "Parrafo 1: Who are you and what specialty interests you?",
            "Parrafo 2: What tools, tasks or systems are important in that specialty?",
            "Parrafo 3: How has that specialty changed? Use used to.",
            "Parrafo 4: What have you learned, liked or noticed? Use Present Perfect with since and for.",
            "Parrafo 5: Why is English useful in this technical field?",
        ],
    )

    add_support_table(
        doc,
        "4. Banco de apoyo para el trabajo escrito",
        [
            (
                "Inicio de texto",
                "Hello, my name is... / I am a first-year student. / The specialty that interests me most is...",
            ),
            (
                "Vocabulario tecnico posible",
                "engine, brake, wrench, lathe, circuit, multimeter, resistor, soldering iron, ink, printer, safety goggles, battery, panel, tool",
            ),
            (
                "Present Perfect",
                "I have liked this specialty for... / I have been interested in... since... / I have learned about... since...",
            ),
            (
                "Used to",
                "Mechanics used to... / Electricians used to... / Printers used to... / Technicians used to...",
            ),
            (
                "Cierre",
                "English is important because... / In the future, I want to... / This specialty is important for society because...",
            ),
        ],
        COLORS["teal"],
    )

    add_rubric_table(
        doc,
        "5. Pauta de evaluacion - Expresion escrita",
        [
            (
                "Cumplimiento de la tarea",
                "Responde completamente la consigna y desarrolla todas las partes solicitadas.",
                "Responde la mayor parte de la consigna con leves omisiones.",
                "Responde parcialmente la consigna y faltan varios elementos.",
                "No responde a la consigna o el texto es muy incompleto.",
            ),
            (
                "Vocabulario tecnico",
                "Usa 5 o mas terminos tecnicos de forma pertinente y clara.",
                "Usa 4 terminos tecnicos pertinentes.",
                "Usa 2 o 3 terminos tecnicos con apoyo o poca precision.",
                "Usa 0 o 1 termino tecnico o los usa incorrectamente.",
            ),
            (
                "Gramatica del semestre",
                "Usa correctamente Present Perfect, since/for y used to.",
                "Usa correctamente dos de las tres estructuras solicitadas.",
                "Usa solo una estructura correctamente o presenta varios errores.",
                "No evidencia manejo de las estructuras trabajadas.",
            ),
            (
                "Organizacion y coherencia",
                "El texto esta bien organizado, con ideas conectadas y faciles de seguir.",
                "El texto es entendible y mayormente ordenado.",
                "La organizacion es irregular y dificulta por momentos la comprension.",
                "Las ideas aparecen desordenadas y cuesta comprender el texto.",
            ),
            (
                "Escritura y correccion formal",
                "Presenta buena ortografia en ingles y redaccion clara.",
                "Presenta algunos errores, pero no afectan mayormente la comprension.",
                "Presenta varios errores que afectan parcialmente la comprension.",
                "Presenta muchos errores que dificultan comprender el mensaje.",
            ),
        ],
    )

    add_box_heading(doc, "6. Trabajo 2 - Expresion oral (20 puntos)", COLORS["blue"])
    add_body_paragraph(doc, "Titulo sugerido: My Future Career and Technical English", bold=True)
    add_body_paragraph(
        doc,
        "Graba un audio o video en ingles de 3 a 4 minutos. Debes hablar sobre la especialidad tecnica que te interesa y explicar por que seria importante para tu futuro. Puedes apoyarte con una tarjeta de notas, pero no debes leer todo el texto.",
    )
    add_body_paragraph(doc, "Tu grabacion debe incluir obligatoriamente:", bold=True)
    add_numbered(
        doc,
        [
            "Presentacion personal breve.",
            "Nombre de la especialidad que te interesa y al menos 2 razones de tu eleccion.",
            "Descripcion de al menos 3 herramientas, tareas o sistemas relacionados con esa especialidad.",
            "Al menos 2 oraciones con Present Perfect, incluyendo since o for.",
            "Al menos 2 oraciones con used to para mostrar cambios entre pasado y presente.",
            "Una conclusion sobre la importancia del ingles en el mundo tecnico-profesional.",
        ],
    )
    add_bullets(
        doc,
        [
            "Habla lento y claro. Es mejor usar oraciones simples y correctas que oraciones muy largas con errores.",
            "Puedes practicar primero leyendo tu guion y luego grabar sin leer completamente.",
            "Si grabas video, no es obligatorio aparecer de cuerpo completo: basta con que se escuche con claridad.",
        ],
    )

    add_support_table(
        doc,
        "7. Banco de apoyo para el trabajo oral",
        [
            (
                "Frases de presentacion",
                "Hello, my name is... / I am in first year of high school. / Today I want to talk about...",
            ),
            (
                "Para explicar intereses",
                "I am interested in... because... / I want to study... because... / This specialty is important because...",
            ),
            (
                "Para hablar del presente",
                "Workers use... / Technicians check... / This specialty includes...",
            ),
            (
                "Para hablar de cambios",
                "Workers used to... but now they... / In the past, technicians used to...",
            ),
            (
                "Para cerrar",
                "English is useful because... / In the future, I want to... / Thank you for listening.",
            ),
        ],
        COLORS["rose"],
    )

    add_rubric_table(
        doc,
        "8. Pauta de evaluacion - Expresion oral",
        [
            (
                "Cumplimiento de la tarea",
                "Desarrolla completamente la presentacion y aborda todos los puntos solicitados.",
                "Desarrolla la mayor parte de la presentacion con pequenas omisiones.",
                "Desarrolla parcialmente la presentacion y faltan varios puntos.",
                "La presentacion es muy incompleta o no responde a la tarea.",
            ),
            (
                "Vocabulario tecnico",
                "Usa 5 o mas terminos tecnicos de forma pertinente y clara.",
                "Usa 4 terminos tecnicos pertinentes.",
                "Usa 2 o 3 terminos tecnicos con apoyo o poca precision.",
                "Usa 0 o 1 termino tecnico o los usa incorrectamente.",
            ),
            (
                "Gramatica del semestre",
                "Usa correctamente Present Perfect, since/for y used to durante la presentacion.",
                "Usa correctamente dos de las tres estructuras solicitadas.",
                "Usa solo una estructura correctamente o presenta errores frecuentes.",
                "No evidencia manejo de las estructuras trabajadas.",
            ),
            (
                "Pronunciacion y claridad",
                "Se entiende claramente durante toda la grabacion.",
                "Se entiende bien en la mayor parte de la grabacion.",
                "A veces cuesta entender algunas palabras o frases.",
                "La pronunciacion dificulta seriamente la comprension.",
            ),
            (
                "Fluidez y seguridad",
                "Habla con ritmo adecuado y poco apoyo de lectura.",
                "Habla con algunas pausas, pero mantiene el mensaje.",
                "Habla con muchas pausas o depende demasiado del guion.",
                "Lee completamente o presenta gran dificultad para sostener el discurso.",
            ),
        ],
    )

    add_box_heading(doc, "9. Criterio de calificacion final", COLORS["blue"])
    add_bullets(
        doc,
        [
            "Trabajo escrito: 20 puntos.",
            "Trabajo oral: 20 puntos.",
            "Puntaje total: 40 puntos.",
            "Exigencia institucional: 60%.",
            "El docente podra retroalimentar ambos productos considerando el avance individual del estudiante y su situacion de continuidad no presencial.",
        ],
    )

    return doc


def build_guide_document():
    doc = create_document()

    add_logo(doc)
    add_title(
        doc,
        "Guia de Apoyo - Evaluacion Especial de Ingles - 1ro Medio",
        "Herramientas, modelos y organizadores para resolver las actividades de forma autonoma",
    )
    add_info_table(
        doc,
        evaluation_name="Guia de apoyo complementaria",
        due_date="Usar junto a la evaluacion",
        total_points="No aplica",
        products="Apoyo para trabajo escrito y oral",
        objective="Asegurar que el estudiante cuente con explicaciones, ejemplos, vocabulario, estructuras y checklist suficientes para desarrollar autonomamente la evaluacion especial del semestre.",
        exigency="No aplica",
    )

    add_box_heading(doc, "1. Como usar esta guia", COLORS["blue"])
    add_bullets(
        doc,
        [
            "Lee primero la evaluacion especial completa y luego vuelve a esta guia para planificar cada producto.",
            "Puedes usar esta guia mientras escribes tu texto y mientras preparas tu audio o video.",
            "No necesitas memorizar todo: la idea es que tengas ejemplos y ayudas concretas para responder.",
            "Si una palabra te cuesta, usa una mas simple. Lo importante es comunicar tus ideas de forma clara y personal.",
        ],
    )

    add_box_heading(doc, "2. Ruta sugerida de trabajo", COLORS["blue"])
    add_numbered(
        doc,
        [
            "Elige una especialidad tecnica que conozcas o que te interese: Automotriz, Industrial, Electricidad, Electronica o Grafica.",
            "Subraya en esta guia el vocabulario y las frases que si podrias usar.",
            "Haz un borrador corto en espanol con tus ideas principales.",
            "Transforma esas ideas al ingles usando los modelos de esta guia.",
            "Escribe primero el trabajo escrito.",
            "Despues usa parte de ese texto como base para preparar el trabajo oral, pero no lo leas completo al grabar.",
            "Revisa el checklist final antes de entregar.",
        ],
    )

    add_support_table(
        doc,
        "3. Vocabulario tecnico base por especialidad",
        [
            ("Mecanica Automotriz", "engine, brake, wrench, battery, oil, scanner, tire, mechanic"),
            ("Mecanica Industrial", "lathe, machine, metal, tool, caliper, blueprint, workshop, process"),
            ("Electricidad", "wire, circuit, breaker, panel, outlet, voltage, multimeter, electrician"),
            ("Electronica", "resistor, capacitor, PCB, signal, sensor, soldering iron, LED, technician"),
            ("Grafica", "ink, printer, paper, CMYK, design, press, cutter, laminator"),
        ],
        COLORS["amber"],
    )

    add_support_table(
        doc,
        "4. Recordatorio gramatical del semestre",
        [
            (
                "Present Perfect",
                "Formula: subject + have/has + past participle. Ejemplos: I have learned technical vocabulary. She has used a multimeter.",
            ),
            (
                "Since",
                "Se usa con un punto de inicio. Ejemplos: since 2024 / since last year / since I was 12.",
            ),
            (
                "For",
                "Se usa con un periodo de tiempo. Ejemplos: for two years / for six months / for a long time.",
            ),
            (
                "Used to",
                "Se usa para hablar del pasado. Ejemplo: Workers used to use simple tools, but now they use digital machines.",
            ),
            (
                "Conectores utiles",
                "because, also, in the past, now, in the future, finally, for example",
            ),
        ],
        COLORS["teal"],
    )

    add_box_heading(doc, "5. Organizador para el trabajo escrito", COLORS["blue"])
    add_body_paragraph(doc, "Usa este esquema antes de escribir tu version final.", bold=True)
    add_support_table(
        doc,
        "5.1 Plan de parrafos",
        [
            ("Parrafo 1", "Presentate y nombra la especialidad que te interesa. Modelo: My name is... I am a first-year student. I am interested in..."),
            ("Parrafo 2", "Describe herramientas, tareas o sistemas. Modelo: In this specialty, workers use... They also check..."),
            ("Parrafo 3", "Compara pasado y presente con used to. Modelo: Technicians used to... but now they..."),
            ("Parrafo 4", "Usa Present Perfect con since y for. Modelo: I have liked this specialty for... / I have been interested in it since..."),
            ("Parrafo 5", "Cierra explicando por que el ingles es importante. Modelo: English is important because..."),
        ],
        COLORS["rose"],
    )

    add_box_heading(doc, "6. Modelo de texto orientador", COLORS["blue"])
    add_body_paragraph(
        doc,
        "Este modelo no debe copiarse literalmente. Sirve para que veas como unir vocabulario, Present Perfect y used to en un solo texto.",
        italic=True,
    )
    add_body_paragraph(
        doc,
        "My name is Tomas and I am a first-year student. I am interested in Electricity because I like circuits, tools and technical work. In this specialty, electricians use wire, panels and multimeters. They check voltage and install electrical systems in buildings. In the past, electricians used to test circuits with simple tools, but now they use digital equipment. They also used to draw plans by hand, but now many workers use software. I have liked technical activities for three years, and I have been interested in Electricity since I visited a workshop with my family. I think English is important because many manuals, warning labels and technical videos are in English. In the future, I want to learn more about this specialty and use English in my studies and work.",
        size=8.5,
    )

    add_box_heading(doc, "7. Organizador para el trabajo oral", COLORS["blue"])
    add_bullets(
        doc,
        [
            "Inicio: Say your name and the topic. Example: Hello, my name is... Today I want to talk about...",
            "Desarrollo 1: Explain why you like the specialty.",
            "Desarrollo 2: Mention 3 tools, tasks or systems.",
            "Desarrollo 3: Compare past and present with used to.",
            "Desarrollo 4: Use Present Perfect with since and for.",
            "Cierre: Explain why English is useful and thank the teacher.",
        ],
    )
    add_support_table(
        doc,
        "7.1 Frases modelo para grabar",
        [
            ("Abrir la grabacion", "Hello, my name is... Today I want to talk about the specialty of..."),
            ("Explicar razones", "I like this specialty because... / It is interesting because..."),
            ("Describir herramientas", "Workers use... / Technicians work with... / One important tool is..."),
            ("Comparar cambios", "In the past, workers used to... but now they..."),
            ("Hablar de experiencia o interes", "I have liked this specialty for... / I have been interested in it since..."),
            ("Cerrar", "English is important because... Thank you for listening."),
        ],
        COLORS["light_blue"],
    )

    add_box_heading(doc, "8. Errores frecuentes que puedes evitar", COLORS["blue"])
    add_bullets(
        doc,
        [
            "No escribas frases sueltas sin relacion. Trata de conectar tus ideas.",
            "No confundas since con for. Since = inicio. For = duracion.",
            "No escribas use to en afirmativo pasado. La forma correcta es used to.",
            "No repitas la misma palabra demasiadas veces si puedes alternar con worker, technician, mechanic, electrician, printer.",
            "No leas demasiado rapido en la grabacion. Hablar claro vale mas que hablar rapido.",
        ],
    )

    add_box_heading(doc, "9. Checklist antes de entregar", COLORS["blue"])
    add_bullets(
        doc,
        [
            "Mi texto escrito tiene entre 180 y 220 palabras, o al menos una extension suficiente y bien desarrollada.",
            "Inclui vocabulario tecnico relacionado con la especialidad.",
            "Use Present Perfect.",
            "Use since y for correctamente.",
            "Use used to para comparar pasado y presente.",
            "Explique por que el ingles es importante en el area tecnica.",
            "Mi audio o video dura aproximadamente entre 3 y 4 minutos.",
            "En la grabacion se entiende mi voz con claridad.",
            "Revise ortografia basica y pronunciacion antes de entregar.",
        ],
    )

    add_box_heading(doc, "10. Mensaje final para el estudiante", COLORS["blue"])
    add_body_paragraph(
        doc,
        "Con esta guia tienes vocabulario, estructuras, ejemplos, modelos de respuesta y un checklist de revision. La idea es que puedas desarrollar ambos trabajos con autonomia y con apoyos suficientes. No se espera perfeccion, sino una respuesta personal, clara y coherente con lo trabajado durante el semestre.",
    )

    return doc


def save_document(doc, output_paths):
    for output_path in output_paths:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"DOCX generado: {output_path}")


def main():
    evaluation_doc = build_document()
    save_document(evaluation_doc, EVALUATION_OUTPUT_PATHS)

    guide_doc = build_guide_document()
    save_document(guide_doc, GUIDE_OUTPUT_PATHS)


if __name__ == "__main__":
    main()