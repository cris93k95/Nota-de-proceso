#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "Profe Jessica" / "Resumen_pilares_reconstruccion_nacional.docx"


SECTIONS = [
    (
        "1. Reconstruccion fisica",
        [
            "Amplia el Fondo de Emergencia Transitorio por Incendios para responder a los siniestros de Nuble y Biobio de 2026 y financiar la reconstruccion de 4.429 viviendas.",
            "Busca asegurar recursos para la emergencia mediante medidas tributarias transitorias, entre ellas la declaracion voluntaria de bienes o rentas en el extranjero con impuesto unico de 10% y una tasa de 7% si esos recursos se repatrian e invierten en Chile.",
            "Permite regularizar saldos historicos acumulados en registros empresariales como FUR, STUT y retiros en exceso del FUT con un impuesto sustitutivo de 10%, para movilizar capital y recaudar fondos.",
            "Autoriza a municipalidades y a la Tesoreria General de la Republica a otorgar alivios por deudas, multas e intereses a personas naturales y mipymes, favoreciendo la formalizacion y la continuidad economica.",
            "Reduce en 50% el impuesto a las donaciones por un tiempo limitado para acelerar transferencias patrimoniales y convertir expectativas de recaudacion futura en ingresos mas inmediatos.",
        ],
    ),
    (
        "2. Reconstruccion economica",
        [
            "Impulsa la construccion y la vivienda con una exencion temporal de IVA para la venta de viviendas nuevas y un nuevo regimen tributario para propiedades DFL 2 de hasta 90 m2 destinadas al arriendo.",
            "Entrega mayor certeza a la inversion al reducir de dos anos a seis meses el plazo para invalidar autorizaciones sectoriales.",
            "Propone reactivar la acuicultura facilitando micro relocalizaciones, reemplazando la caducidad por no uso por un cobro de patente incrementada y agilizando informes tecnicos y ambientales.",
            "Crea un credito tributario para proteger el empleo formal, con mayor apoyo para remuneraciones mas bajas, buscando reducir la informalidad y el costo de contratacion.",
            "Rebaja gradualmente el impuesto de primera categoria hasta 23%, restablece la integracion total del sistema tributario y elimina el impuesto unico de 10% a ciertas ganancias de capital bursatiles.",
            "Establece invariabilidad tributaria por 25 anos para inversiones iguales o superiores a 50 millones de dolares, con el objetivo de recuperar certeza juridica y atraer capitales.",
            "Exime del pago de contribuciones a personas de 65 anos o mas respecto de su vivienda principal, como alivio patrimonial a adultos mayores.",
        ],
    ),
    (
        "3. Reconstruccion institucional",
        [
            "Busca racionalizar la llamada permisologia, haciendo mas proporcionadas las exigencias regulatorias y acelerando procesos clave para la inversion.",
            "En evaluacion ambiental, propone un regimen con menos adendas, mayor rectoria tecnica del SEA y limites a requerimientos de organismos que excedan sus competencias.",
            "Refuerza la certeza juridica de las RCA al impedir su invalidacion administrativa por la via general, concentrando su impugnacion en los mecanismos propios de la legislacion ambiental.",
            "Crea un sistema de restitucion de gastos para titulares cuyas RCA favorables sean anuladas judicialmente, reconociendo la confianza legitima generada por la autorizacion estatal.",
            "Limita la paralizacion de proyectos mediante medidas cautelares: fija plazos breves, tope maximo de seis meses y posibilidad de apelacion ante Corte de Apelaciones.",
            "Simplifica la gestion del Consejo de Monumentos Nacionales con plazos fatales, certificacion electronica por silencio y una categoria de intervenciones menores para obras de bajo impacto.",
        ],
    ),
    (
        "4. Reconstruccion fiscal",
        [
            "Plantea contener el crecimiento del gasto publico, especialmente en gratuidad de educacion superior, sin afectar a quienes ya reciben el beneficio.",
            "Congela por dos anos el ingreso de nuevas instituciones al regimen de financiamiento institucional de gratuidad y modifica las condiciones para expandir la cobertura futura.",
            "Elimina la franquicia tributaria del SENCE por considerar insuficiente su evidencia de impacto frente a su alto costo fiscal.",
            "Aumenta los cupos del incentivo al retiro en universidades estatales y administracion central, buscando ordenar dotaciones y mejorar eficiencia del gasto.",
            "Endurece las sanciones por uso indebido de licencias medicas en el Estado, calificandolo como una infraccion grave al principio de probidad.",
            "Fortalece la fiscalizacion y la eficiencia del gasto mediante mayor cruce de informacion entre el SII, Desarrollo Social y la DIPRES, con foco en mejor focalizacion y control.",
        ],
    ),
]


def set_normal_style(document):
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)


def add_title(document, title, subtitle):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1F3A5F")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(10)


def add_paragraph(document, text, *, bold=False, italic=False, color=None, space_after=4):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_section(document, heading, bullets):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(heading)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("1F3A5F")

    for bullet in bullets:
        item = document.add_paragraph(style="List Bullet")
        item.paragraph_format.left_indent = Cm(0.7)
        item.paragraph_format.space_after = Pt(2)
        run = item.add_run(bullet)
        run.font.size = Pt(10.5)


def build_document():
    document = Document()
    set_normal_style(document)

    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(
        document,
        "Resumen de los principales pilares del proyecto de ley para la Reconstruccion Nacional",
        "Basado en el Mensaje Presidencial N° 018-374, de 22 de abril de 2026",
    )

    add_paragraph(
        document,
        "Nota importante: el texto fuente corresponde a un proyecto de ley (mensaje presidencial), por lo que este resumen describe sus ejes y medidas propuestas, no una ley ya promulgada.",
        italic=True,
        color="7A4E00",
        space_after=8,
    )

    add_paragraph(
        document,
        "Objetivo general: reactivar el crecimiento economico, acelerar la reconstruccion tras los incendios, recuperar la inversion, crear empleo formal, racionalizar permisos y fortalecer la sostenibilidad fiscal del Estado.",
        bold=True,
        space_after=8,
    )

    add_paragraph(
        document,
        "El mensaje organiza su propuesta en cuatro ejes complementarios: reconstruccion fisica, reconstruccion economica, reconstruccion institucional y reconstruccion fiscal.",
        space_after=6,
    )

    for heading, bullets in SECTIONS:
        add_section(document, heading, bullets)

    add_section(
        document,
        "Resultados esperados segun el mensaje",
        [
            "Creacion aproximada de 180 mil empleos en el corto plazo por reactivacion de sectores paralizados y mayor ingreso de capitales.",
            "Reduccion de la tasa de desempleo desde 8,3% a 6,5% hacia 2030.",
            "Recuperacion de un crecimiento promedio cercano a 4% anual y restablecimiento del equilibrio fiscal estructural.",
            "Mejora en certeza juridica e incentivos para proyectos de inversion, vivienda, infraestructura y desarrollo productivo.",
        ],
    )

    add_paragraph(
        document,
        "En sintesis, el proyecto combina medidas de emergencia y reconstruccion con una agenda de crecimiento, desregulacion, alivio tributario, fortalecimiento de la inversion y contencion fiscal.",
        bold=True,
        space_after=0,
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(f"Documento generado en: {OUTPUT_PATH}")