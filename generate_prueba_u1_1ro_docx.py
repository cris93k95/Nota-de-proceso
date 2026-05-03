#!/usr/bin/env python3
"""
Genera la Prueba de Comprensión Lectora U1 1ro Medio en formato Word (.docx)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from generate_prueba_u1_1ro import (
    GENERAL_INSTRUCTIONS,
    GRAMMAR_ACTIVITIES,
    HEADER_SUMMARY_ITEMS,
    LOGO_PATH,
    PASSING_PERCENTAGE,
    READING_TEXTS,
    TEST_OBJECTIVE,
    TEST_SKILLS,
    TEST_TITLE,
    TOTAL_POINTS,
    VOCAB_MATCHING,
)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "PLANIFICACIONES_2026_LISTO_IMPRESION")

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_part_header(doc, text):
    """Add a blue part header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Blue background via shading
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): '2962FF',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)
    return p

def add_instructions(doc, text):
    """Add italic instructions."""
    clean_text = text.replace("<br/>", "\n").replace("<br>", "\n")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(clean_text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_reading_text(doc, title, text):
    """Add a reading text block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    # Shading for reading text background
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F5',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)
    title_run = p.add_run(title + "\n")
    title_run.bold = True
    title_run.font.size = Pt(9)
    text_run = p.add_run(text)
    text_run.font.size = Pt(8.5)
    return p

def add_question(doc, num, q_text):
    """Add a question."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{num}. {q_text}")
    run.bold = True
    run.font.size = Pt(8.5)
    return p

def add_option(doc, text):
    """Add an option."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(8)
    return p


def add_header_panel(doc):
    if LOGO_PATH and os.path.exists(LOGO_PATH):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(LOGO_PATH, width=Cm(17.99))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run(TEST_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(11)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_before = Pt(1)
    summary.paragraph_format.space_after = Pt(3)
    for index, (label, value) in enumerate(HEADER_SUMMARY_ITEMS):
        if index:
            separator = summary.add_run("   ")
            separator.font.size = Pt(8.2)
        label_run = summary.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(8.2)
        value_run = summary.add_run(str(value))
        value_run.font.size = Pt(8.2)

    student_line = doc.add_paragraph()
    student_line.paragraph_format.space_before = Pt(1)
    student_line.paragraph_format.space_after = Pt(3)
    for label, value in (
        ("Nombre:", " ________________________________"),
        ("   Curso:", " __________"),
        ("   Fecha:", " ___/___/2026"),
    ):
        label_run = student_line.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(8.2)
        value_run = student_line.add_run(value)
        value_run.font.size = Pt(8.2)

    objective = doc.add_paragraph()
    objective.paragraph_format.space_before = Pt(1)
    objective.paragraph_format.space_after = Pt(1)
    objective_label = objective.add_run("Objetivo: ")
    objective_label.bold = True
    objective_label.font.size = Pt(8)
    objective_text = objective.add_run(TEST_OBJECTIVE)
    objective_text.font.size = Pt(8)

    skills = doc.add_paragraph()
    skills.paragraph_format.space_before = Pt(0)
    skills.paragraph_format.space_after = Pt(3)
    skills_label = skills.add_run("Habilidades: ")
    skills_label.bold = True
    skills_label.font.size = Pt(8)
    skills_text = skills.add_run(TEST_SKILLS)
    skills_text.font.size = Pt(8)

    instructions_title = doc.add_paragraph()
    instructions_title.paragraph_format.space_before = Pt(2)
    instructions_title.paragraph_format.space_after = Pt(1)
    title_run = instructions_title.add_run("Instrucciones generales")
    title_run.bold = True
    title_run.font.size = Pt(8)

    for item in GENERAL_INSTRUCTIONS:
        instruction = doc.add_paragraph()
        instruction.paragraph_format.space_before = Pt(0)
        instruction.paragraph_format.space_after = Pt(0)
        instruction.paragraph_format.left_indent = Cm(0.35)
        run = instruction.add_run(f"- {item}")
        run.font.size = Pt(7.8)


def strip_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = tcBorders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none',
                    qn('w:sz'): '0',
                    qn('w:space'): '0',
                    qn('w:color'): 'auto'
                })
                tcBorders.append(border)
            tcPr.append(tcBorders)


def collect_answers():
    answers = {}
    for text in READING_TEXTS:
        for question in text["questions"]:
            answers[question["num"]] = question["answer"]
    for number, answer in VOCAB_MATCHING["answers"].items():
        answers[number] = answer
    for activity in GRAMMAR_ACTIVITIES:
        for question in activity["questions"]:
            answers[question["num"]] = question["answer"]
    return answers

def generate_docx():
    doc = Document()

    # Page setup - Letter size
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    add_header_panel(doc)

    # ============================================================
    # PART I: READING COMPREHENSION
    # ============================================================
    add_part_header(doc, "PART I — READING COMPREHENSION (15 points)")
    add_instructions(doc, "Questions 1–15\nRead each text carefully and choose the correct answer for each question.\nCircle the letter of your answer (A, B, C, D or E).")

    for text in READING_TEXTS:
        add_reading_text(doc, text["title"], text["text"])
        for question in text["questions"]:
            add_question(doc, question["num"], question["q"])
            for opt in question["options"]:
                add_option(doc, opt)

    # ============================================================
    # PART II: VOCABULARY MATCHING
    # ============================================================
    add_part_header(doc, "PART II — VOCABULARY: Matching / Términos Pareados (10 points)")
    add_instructions(doc, VOCAB_MATCHING["instructions"])

    # Create matching table
    terms = VOCAB_MATCHING["terms"]
    definitions = VOCAB_MATCHING["definitions"]

    match_table = doc.add_table(rows=max(len(terms), len(definitions)), cols=4)
    match_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(max(len(terms), len(definitions))):
        row = match_table.rows[i]
        if i < len(terms):
            row.cells[0].text = f"{terms[i]['num']}."
            row.cells[1].text = f"{terms[i]['term']}  _____"
            for cell_idx in [0, 1]:
                for paragraph in row.cells[cell_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
        if i < len(definitions):
            row.cells[2].text = f"{definitions[i]['letter']})"
            row.cells[3].text = definitions[i]["definition"]
            for cell_idx in [2, 3]:
                for paragraph in row.cells[cell_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)

    strip_table_borders(match_table)

    # ============================================================
    # PART III: GRAMMAR
    # ============================================================
    add_part_header(doc, "PART III — GRAMMAR (15 points)")
    add_instructions(doc, "Questions 26–40\nUse the information from Part I to complete each sentence correctly.\nCircle the letter of your answer (A, B, C, D or E).")

    for activity in GRAMMAR_ACTIVITIES:
        add_part_header(doc, activity["title"])
        add_instructions(doc, activity["instructions"])
        for question in activity["questions"]:
            add_question(doc, question["num"], question["q"])
            for opt in question["options"]:
                add_option(doc, opt)

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("— Fin de la evaluación — Revisa tus respuestas antes de entregar. —")
    run.italic = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ============================================================
    # ANSWER KEY (new page)
    # ============================================================
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pauta de Respuestas Correctas")
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TEST_TITLE)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    answers = collect_answers()

    # Answer grid table (8 rows x 5 cols)
    ans_table = doc.add_table(rows=8, cols=5)
    ans_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sorted_numbers = sorted(answers)
    for i, num in enumerate(sorted_numbers):
        row_idx = i % 8
        col_idx = i // 8
        cell = ans_table.cell(row_idx, col_idx)
        cell.text = f"{num}. {answers[num]}"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

    strip_table_borders(ans_table)

    # Score conversion table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Tabla de Conversión de Puntaje a Nota ({PASSING_PERCENTAGE}% de exigencia)")
    run.bold = True
    run.font.size = Pt(8)

    total_points = TOTAL_POINTS
    passing = total_points * (PASSING_PERCENTAGE / 100)

    score_table = doc.add_table(rows=1, cols=6)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_table.style = 'Table Grid'
    headers = ["Puntaje", "Nota", "Puntaje", "Nota", "Puntaje", "Nota"]
    for i, h in enumerate(headers):
        cell = score_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'E8E8E8')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)

    scores = []
    for pts in range(0, total_points + 1):
        if pts < passing:
            nota = 1.0 + 3.0 * (pts / passing)
        else:
            nota = 4.0 + 3.0 * ((pts - passing) / (total_points - passing))
        nota = min(nota, 7.0)
        nota = round(nota, 1)
        scores.append((pts, nota))

    rows_per_col = (len(scores) + 2) // 3
    for i in range(rows_per_col):
        row = score_table.add_row()
        for col in range(3):
            idx = i + col * rows_per_col
            if idx < len(scores):
                pts, nota = scores[idx]
                pts_cell = row.cells[col * 2]
                nota_cell = row.cells[col * 2 + 1]
                pts_cell.text = str(pts)
                nota_cell.text = str(nota)
                if nota >= 4.0:
                    set_cell_shading(nota_cell, 'E8F5E9')
                for c in [pts_cell, nota_cell]:
                    for paragraph in c.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(7)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, "PRUEBA_COMPRENSION_LECTORA_U1_1RO_MEDIO.docx")
    doc.save(output_path)
    print(f"DOCX generado: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_docx()
