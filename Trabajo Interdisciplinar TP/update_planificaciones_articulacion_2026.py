import argparse
import json
from copy import deepcopy
from pathlib import Path

from docx import Document

from generate_planificaciones_articulacion_2026 import (
    OUTPUT_DIR,
    SPECIALTY_LETTERS,
    TEMPLATE_PATH,
    base_1ro_payload,
    fill_action_table,
    fill_articulation_table,
    fill_calendar_table,
    fill_project_table,
    fill_version_table,
    payload_3ro,
    payload_4to,
)


LETTER_TO_SPECIALTY = {value: key for key, value in SPECIALTY_LETTERS.items()}


def deep_update(target, overrides):
    for key, value in overrides.items():
        if isinstance(value, dict) and isinstance(target.get(key), dict):
            deep_update(target[key], value)
        else:
            target[key] = value
    return target


def parse_doc_name(doc_name):
    if doc_name == "1ro Medio - Ingles articulacion TP.docx":
        return "1ro", None

    stem = Path(doc_name).stem
    if " - " not in stem:
        raise ValueError(f"No se pudo interpretar el nombre del archivo: {doc_name}")

    course_part, specialty = stem.split(" - ", 1)
    if course_part.startswith("3ro "):
        return "3ro", specialty
    if course_part.startswith("4to "):
        return "4to", specialty

    raise ValueError(f"Curso no soportado en el archivo: {doc_name}")


def build_payload(level, specialty=None):
    if level == "1ro":
        return base_1ro_payload()
    if level == "3ro" and specialty:
        return payload_3ro(specialty)
    if level == "4to" and specialty:
        return payload_4to(specialty)
    raise ValueError("Combinación de curso/especialidad no válida")


def update_document(doc_path, payload):
    source_path = doc_path if doc_path.exists() else TEMPLATE_PATH
    doc = Document(source_path)
    fill_version_table(doc.tables[0], payload["version_description"])
    fill_project_table(doc.tables[2], payload)
    fill_articulation_table(doc.tables[3], payload)
    fill_action_table(doc.tables[4], payload)
    fill_calendar_table(doc.tables[5], payload["calendar_rows"])
    doc.save(doc_path)


def load_json(path):
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def update_one(args):
    level = args.level
    specialty = args.specialty
    if args.doc_name:
        level, specialty = parse_doc_name(args.doc_name)

    payload = deepcopy(build_payload(level, specialty))
    if args.overrides:
        deep_update(payload, load_json(args.overrides))

    if args.output:
        output_path = Path(args.output)
    elif args.doc_name:
        output_path = OUTPUT_DIR / args.doc_name
    elif level == "1ro":
        output_path = OUTPUT_DIR / "1ro Medio - Ingles articulacion TP.docx"
    else:
        output_path = OUTPUT_DIR / f"{level} {SPECIALTY_LETTERS[specialty]} - {specialty}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    update_document(output_path, payload)
    print(f"Actualizado: {output_path}")


def update_batch(batch_file):
    batch_data = load_json(batch_file)
    for doc_name, overrides in batch_data.items():
        level, specialty = parse_doc_name(doc_name)
        payload = deepcopy(build_payload(level, specialty))
        deep_update(payload, overrides)
        output_path = OUTPUT_DIR / doc_name
        update_document(output_path, payload)
        print(f"Actualizado: {output_path}")


def build_parser():
    parser = argparse.ArgumentParser(
        description="Actualiza planificaciones de articulación 2026 a partir de la plantilla base y overrides JSON."
    )
    parser.add_argument("--batch", help="Ruta a un JSON con múltiples documentos a actualizar.")
    parser.add_argument("--doc-name", help="Nombre del DOCX dentro de la carpeta de salida.")
    parser.add_argument("--level", choices=["1ro", "3ro", "4to"], help="Curso a actualizar.")
    parser.add_argument("--specialty", help="Especialidad del curso 3ro/4to.")
    parser.add_argument("--overrides", help="Ruta a un JSON con campos a sobreescribir para un documento.")
    parser.add_argument("--output", help="Ruta de salida opcional. Si no se entrega, actualiza el archivo estándar.")
    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    if args.batch:
        update_batch(args.batch)
        return

    if not args.doc_name and not args.level:
        parser.error("Debes indicar --batch o bien --doc-name / --level.")

    if args.level in {"3ro", "4to"} and not args.specialty and not args.doc_name:
        parser.error("Para 3ro y 4to debes indicar --specialty o usar --doc-name.")

    update_one(args)


if __name__ == "__main__":
    main()