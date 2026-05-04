from copy import deepcopy
from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "Formato planificación articulación 2026 (1).docx"
OUTPUT_DIR = BASE_DIR / "Planificaciones articulacion 2026"
TODAY = "21-04-2026"


SPECIALTIES = [
    "Mecánica Industrial",
    "Mecánica Automotriz",
    "Electricidad",
    "Gráfica",
    "Electrónica",
]

SPECIALTY_LETTERS = {
    "Mecánica Industrial": "A",
    "Mecánica Automotriz": "B",
    "Electricidad": "C",
    "Gráfica": "D",
    "Electrónica": "E",
}


def set_cell(table, row_idx, col_idx, text):
    table.rows[row_idx].cells[col_idx].text = text


def fill_version_table(table, description):
    set_cell(table, 1, 0, "01")
    set_cell(table, 1, 1, TODAY)
    set_cell(table, 1, 2, description)
    set_cell(table, 1, 3, "Docente de Inglés")
    set_cell(table, 1, 4, "Coordinación Pedagógica HC-TP")


def fill_project_table(table, data):
    fields = [
        data["project_name"],
        data["general_description"],
        data["general_objectives"],
        data["specific_objectives"],
        data["specialties"],
        data["subsectors"],
        data["expected_results"],
        data["methodology"],
    ]
    for idx, value in enumerate(fields):
        set_cell(table, idx, 1, value)


def fill_articulation_table(table, data):
    for col in range(3):
        set_cell(table, 2, col, f"Asignatura/Subsector: {data['general_subject']}")
    for col in range(3, 6):
        set_cell(table, 2, col, f"Módulo Especialidad: {data['specialty_module']}")

    for row_offset, articulation_row in enumerate(data["articulation_rows"], start=4):
        for col_idx, value in enumerate(articulation_row):
            set_cell(table, row_offset, col_idx, value)


def fill_action_table(table, data):
    values = [
        data["key_action_name"],
        data["key_action_objective"],
        data["participants"],
        data["key_action_description"],
        data["achievement_indicators"],
        data["verification"],
    ]
    for row_idx, value in enumerate(values):
        set_cell(table, row_idx, 1, value)


def fill_calendar_table(table, rows):
    month_index = {
        "Mar.": 1,
        "Abr.": 2,
        "May.": 3,
        "Jun.": 4,
        "Jul.": 5,
        "Ago.": 6,
        "Sep.": 7,
        "Oct.": 8,
        "Nov.": 9,
        "Dic.": 10,
    }
    for row_idx, row_data in enumerate(rows, start=1):
        set_cell(table, row_idx, 0, row_data["name"])
        for month in row_data.get("months", []):
            set_cell(table, row_idx, month_index[month], "X")


def build_doc(output_name, payload):
    doc = Document(TEMPLATE_PATH)
    fill_version_table(doc.tables[0], payload["version_description"])
    fill_project_table(doc.tables[2], payload)
    fill_articulation_table(doc.tables[3], payload)
    fill_action_table(doc.tables[4], payload)
    fill_calendar_table(doc.tables[5], payload["calendar_rows"])
    doc.save(OUTPUT_DIR / output_name)


def base_1ro_payload():
    return {
        "version_description": "Creación de planificación de articulación 2026 para 1° Medio en Inglés TP.",
        "project_name": "Exploración temprana del área TP desde Inglés en 1° Medio",
        "general_description": (
            "Durante el primer semestre, la asignatura de Inglés aborda una exploración temprana del área técnico-profesional "
            "por medio de vocabulario, textos y actividades contextualizadas en las cinco especialidades impartidas por el colegio. "
            "El proyecto busca que las y los estudiantes conozcan rasgos básicos de cada especialidad antes de la definición futura "
            "de trayectoria, articulando contenidos gramaticales del currículum de Inglés con contextos TP significativos."
        ),
        "general_objectives": (
            "Integrar contenidos del currículum de Inglés del primer semestre con una aproximación temprana al mundo TP, "
            "favoreciendo la exploración vocacional y la comunicación oral y escrita en contextos técnicos iniciales."
        ),
        "specific_objectives": (
            "1. Reconocer y utilizar vocabulario básico asociado a Electricidad, Mecánica Automotriz, Electrónica, Gráfica y Mecánica Industrial.\n"
            "2. Elaborar una presentación oral en inglés sobre la especialidad que cada estudiante elegiría, justificando su elección, herramientas y evolución histórica durante el siglo XX.\n"
            "3. Resolver una evaluación de comprensión lectora en contexto TP aplicando contenidos gramaticales del currículum.\n"
            "4. Desarrollar una tercera evaluación oral en junio vinculada a la descripción de un proceso o mecanismo técnico usando estructuras trabajadas en clases."
        ),
        "specialties": ", ".join(SPECIALTIES),
        "subsectors": "Inglés",
        "expected_results": (
            "- Presentación oral en inglés sobre una especialidad elegida por cada estudiante.\n"
            "- Evaluación escrita de comprensión lectora contextualizada en el área TP.\n"
            "- Interrogación oral de junio sobre descripción de un proceso o mecanismo técnico.\n"
            "- Mayor familiarización con el vocabulario inicial de las cinco especialidades."
        ),
        "methodology": (
            "Enfoque comunicativo, aprendizaje situado y articulación curricular. Se combinan lectura guiada, desarrollo de vocabulario, "
            "presentaciones orales, interrogación oral y evaluación escrita, utilizando el contexto TP como soporte para los aprendizajes de Inglés."
        ),
        "general_subject": "Inglés",
        "specialty_module": "Articulación transversal con las 5 especialidades TP del establecimiento",
        "articulation_rows": [
            [
                "Exploración vocacional temprana y vocabulario técnico inicial.",
                "Comprender ideas generales y vocabulario específico en textos orales y escritos breves relacionados con contextos técnicos escolares.",
                "Participar con disposición, curiosidad y respeto frente a distintas trayectorias formativas.",
                "Vinculación introductoria con tareas, herramientas y campos ocupacionales de las cinco especialidades.",
                "Reconoce funciones, herramientas y rasgos distintivos básicos de cada especialidad ofrecida por el colegio.",
                "A) Comunicarse oralmente y por escrito con claridad.\nB) Trabajar colaborativamente y con responsabilidad.",
            ],
            [
                "Presentación oral sobre especialidad elegida.",
                "Expresar información, opiniones y justificaciones simples en inglés sobre intereses personales y proyección formativa.",
                "Manifestar iniciativa, perseverancia y escucha activa.",
                "Relación entre la asignatura y los saberes iniciales del área TP seleccionada por el estudiante.",
                "Explica razones de elección, herramientas de uso frecuente y referencias históricas básicas de la especialidad durante el siglo XX.",
                "C) Resolver problemas y comunicar soluciones usando recursos pertinentes.",
            ],
            [
                "Comprensión lectora y evaluación oral de junio.",
                "Aplicar contenidos gramaticales del currículum mediante lectura contextualizada y producción oral guiada.",
                "Trabajar con autonomía, atención y disposición a mejorar.",
                "Descripción de procesos o mecanismos técnicos en lenguaje accesible para estudiantes de 1° Medio.",
                "Lee textos breves en contexto TP y describe oralmente un proceso o mecanismo técnico usando al menos cuatro oraciones en voz pasiva y vocabulario pertinente.",
                "D) Aprender de manera autónoma y utilizar información de forma organizada.",
            ],
        ],
        "key_action_name": "Secuencia de exploración TP en Inglés para 1° Medio",
        "key_action_objective": "Vincular el currículum de Inglés del primer semestre con la exploración inicial de las cinco especialidades técnico-profesionales del establecimiento.",
        "participants": "Docente de Inglés, estudiantes de 1° Medio y coordinación pedagógica HC-TP.",
        "key_action_description": (
            "La secuencia considera clases con vocabulario y lectura en contexto TP, una presentación oral en inglés sobre la especialidad "
            "que cada estudiante elegiría, una evaluación escrita de comprensión lectora en mayo y una tercera evaluación oral en junio "
            "sobre la descripción de un proceso o mecanismo técnico utilizando estructuras gramaticales trabajadas en clases."
        ),
        "achievement_indicators": (
            "Reconoce vocabulario de las cinco especialidades; justifica oralmente la elección de una especialidad; comprende textos TP adaptados; "
            "describe un proceso o mecanismo técnico con apoyo de estructuras gramaticales del semestre."
        ),
        "verification": "Rúbrica de presentación oral, prueba escrita de comprensión lectora, interrogación oral de junio y productos de clase.",
        "calendar_rows": [
            {"name": "Planificación conjunta y selección de focos TP", "months": ["Mar."]},
            {"name": "Trabajo oral 1: especialidad elegida", "months": ["Abr."]},
            {"name": "Evaluación escrita de comprensión lectora TP", "months": ["May."]},
            {"name": "Evaluación oral 3: proceso o mecanismo técnico", "months": ["Jun."]},
        ],
    }


def payload_3ro(specialty):
    module_map = {
        "Electricidad": "Exploración temprana y nivelación de conocimientos en Electricidad",
        "Mecánica Automotriz": "Exploración temprana y nivelación de conocimientos en Mecánica Automotriz",
        "Electrónica": "Exploración temprana y nivelación de conocimientos en Electrónica",
        "Gráfica": "Exploración temprana y nivelación de conocimientos en Gráfica",
        "Mecánica Industrial": "Exploración temprana y nivelación de conocimientos en Mecánica Industrial",
    }
    topic_focus = {
        "Electricidad": "energías renovables, la transición hacia ellas y el hidrógeno verde",
        "Mecánica Automotriz": "confort y seguridad del vehículo, sistemas pasivos y activos, diagnóstico básico de fallas, piezas fijas y móviles del motor, ciclo de Carnot, lubricación, refrigeración y sistema de cuatro tiempos, en una versión simplificada y pertinente al nivel del curso",
        "Electrónica": "electrónica digital y sus componentes clave, domótica, automatización, mecatrónica y robótica, abordadas mediante estudios de caso",
        "Gráfica": 'el software "Fiery Command WorkStation"',
        "Mecánica Industrial": "conceptos introductorios y rutinas laborales vinculadas a la especialidad",
    }
    yearly_scope = {
        "Electricidad": "Estos temas se revisarán a lo largo del año por medio de comprensión lectora, integrando además otros trabajos de la asignatura.",
        "Mecánica Automotriz": "Estos temas se revisarán a lo largo del año por medio de comprensión lectora, articulando además otros trabajos propios de la asignatura.",
        "Electrónica": "Estos temas se revisarán a lo largo del año por medio de comprensión lectora, integrando además otros trabajos y actividades de aplicación.",
        "Gráfica": "Este foco se revisará a lo largo del año por medio de comprensión lectora, análisis de vocabulario técnico y otros trabajos de la asignatura.",
        "Mecánica Industrial": "Los contenidos se desarrollarán mediante comprensión lectora y otras actividades articuladas durante el año escolar.",
    }
    calendar_rows = {
        "Electricidad": [
            {"name": "Planificación conjunta con especialidad", "months": ["Mar."]},
            {"name": "Trabajo oral: habilidades y conocimientos previos", "months": ["Abr."]},
            {"name": "Lectura guiada sobre energías renovables e hidrógeno verde", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión técnica", "months": ["Jun.", "Dic."]},
        ],
        "Mecánica Automotriz": [
            {"name": "Planificación conjunta con especialidad", "months": ["Mar."]},
            {"name": "Trabajo oral: habilidades y conocimientos previos", "months": ["Abr."]},
            {"name": "Lectura guiada sobre seguridad, motor y sistemas básicos", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión técnica", "months": ["Jun.", "Dic."]},
        ],
        "Electrónica": [
            {"name": "Planificación conjunta con especialidad", "months": ["Mar."]},
            {"name": "Trabajo oral: habilidades y conocimientos previos", "months": ["Abr."]},
            {"name": "Lectura guiada y estudios de caso de electrónica y automatización", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión técnica", "months": ["Jun.", "Dic."]},
        ],
        "Gráfica": [
            {"name": "Planificación conjunta con especialidad", "months": ["Mar."]},
            {"name": "Trabajo oral: habilidades y conocimientos previos", "months": ["Abr."]},
            {"name": "Lectura guiada sobre Fiery Command WorkStation", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión técnica", "months": ["Jun.", "Dic."]},
        ],
        "Mecánica Industrial": [
            {"name": "Planificación conjunta con especialidad", "months": ["Mar."]},
            {"name": "Trabajo oral: habilidades y conocimientos previos", "months": ["Abr."]},
            {"name": "Lectura guiada y nivelación de conceptos", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión técnica", "months": ["Jun.", "Dic."]},
        ],
    }
    return {
        "version_description": f"Creación de planificación de articulación 2026 para 3° Medio - {specialty}.",
        "project_name": f"Inglés técnico aplicado a {specialty} en 3° Medio",
        "general_description": (
            f"En 3° Medio, las clases de Inglés abordan únicamente temas asociados a {specialty}, con un foco de exploración temprana y "
            f"nivelación de conocimientos. El trabajo articula el desarrollo del idioma con contenidos mínimos requeridos de la especialidad, "
            f"a partir de vocabulario, textos auténticos adaptados y actividades de reflexión sobre el rol técnico de la especialidad en la sociedad moderna. "
            f"En particular, se trabajará sobre {topic_focus[specialty]}. {yearly_scope[specialty]}"
        ),
        "general_objectives": (
            f"Fortalecer el uso de Inglés en contextos propios de {specialty}, integrando comprensión de textos auténticos adaptados, vocabulario técnico "
            f"y producción oral reflexiva sobre habilidades, conocimientos previos y relevancia social de la especialidad, incorporando progresivamente los focos temáticos propios del área."
        ),
        "specific_objectives": (
            f"1. Reconocer y utilizar vocabulario técnico básico vinculado a {specialty}.\n"
            f"2. Presentar oralmente, en inglés, habilidades y conocimientos previos relacionados con {specialty}.\n"
            f"3. Reflexionar sobre la importancia de {specialty} para la sociedad moderna.\n"
            f"4. Resolver evaluaciones y actividades de comprensión lectora basadas en textos auténticos adaptados y conocimientos mínimos requeridos de la especialidad, consultados con docentes del área.\n"
            f"5. Integrar progresivamente el trabajo sobre {topic_focus[specialty]} a lo largo del año escolar."
        ),
        "specialties": specialty,
        "subsectors": "Inglés",
        "expected_results": (
            f"- Presentación oral en inglés sobre habilidades y conocimientos vinculados a {specialty}.\n"
            f"- Evaluación escrita de comprensión de textos auténticos adaptados al nivel de Inglés del curso.\n"
            f"- Nivelación inicial de vocabulario y conceptos mínimos requeridos de {specialty}.\n"
            f"- Desarrollo progresivo de comprensión lectora sobre {topic_focus[specialty]}."
        ),
        "methodology": (
            "Enfoque comunicativo y aprendizaje situado. Las clases consideran exploración de imágenes y textos del campo profesional, nivelación de "
            f"conceptos técnicos básicos, producción oral guiada y evaluación escrita a partir de documentos y textos adaptados consultados con docentes de especialidad. {yearly_scope[specialty]}"
        ),
        "general_subject": "Inglés",
        "specialty_module": module_map[specialty],
        "articulation_rows": [
            [
                "Vocabulario técnico y exploración del campo profesional.",
                f"Identificar vocabulario técnico básico en inglés relacionado con {specialty} mediante la exploración de textos e imágenes del campo profesional.",
                "Trabajar con curiosidad, respeto y disposición al aprendizaje técnico.",
                f"Introducción a conceptos mínimos requeridos, rutinas laborales y focos como {topic_focus[specialty]}.",
                f"Reconoce conceptos, herramientas, procesos y vocabulario inicial propios de {specialty}.",
                "A) Comunicarse con claridad.\nB) Trabajar colaborativamente.",
            ],
            [
                "Presentación oral sobre habilidades para la especialidad.",
                "Expresar información, opiniones y experiencias previas en inglés con apoyo de vocabulario técnico.",
                "Desarrollar confianza, autoevaluación y escucha activa.",
                f"Reflexión sobre habilidades, conocimientos previos y proyección formativa en {specialty}.",
                f"Expone oralmente habilidades y saberes previos, justificando la importancia de {specialty} para la sociedad moderna.",
                "C) Resolver situaciones comunicativas con recursos adecuados.",
            ],
            [
                "Comprensión de textos auténticos adaptados.",
                "Comprender ideas centrales, datos específicos e inferencias simples en textos técnicos adaptados al nivel del curso.",
                "Trabajar con autonomía y rigurosidad en la interpretación de información.",
                f"Lectura de documentos y textos basados en conocimientos mínimos requeridos de {specialty}, consultados con docentes del área, con foco en {topic_focus[specialty]}.",
                f"Responde adecuadamente a evaluaciones y actividades de comprensión lectora basadas en textos auténticos adaptados sobre {specialty}.",
                "D) Gestionar información y aprender de manera autónoma.",
            ],
        ],
        "key_action_name": f"Secuencia de articulación en Inglés para {specialty}",
        "key_action_objective": f"Relacionar el aprendizaje del Inglés con la exploración temprana y la nivelación de conocimientos en {specialty}.",
        "participants": f"Docente de Inglés, docente(s) de {specialty}, estudiantes de 3° Medio TP y coordinación pedagógica HC-TP.",
        "key_action_description": (
            f"La secuencia contempla clases centradas exclusivamente en {specialty}, una evaluación oral sobre habilidades y conocimientos previos, "
            f"y una evaluación escrita de comprensión de textos auténticos adaptados, construida a partir de contenidos mínimos requeridos del área y "
            f"consultada con los docentes de especialidad. Además, se incorporará de manera progresiva el trabajo sobre {topic_focus[specialty]} mediante comprensión lectora y otros trabajos intermedios durante el año."
        ),
        "achievement_indicators": (
            f"Reconoce vocabulario técnico de {specialty}; presenta oralmente habilidades y saberes previos; interpreta textos auténticos adaptados; "
            f"explica la relevancia social de la especialidad en inglés; y progresa en la comprensión de textos asociados a {topic_focus[specialty]}."
        ),
        "verification": "Rúbrica de presentación oral, prueba escrita de comprensión lectora, guías de trabajo y productos de clase.",
        "calendar_rows": calendar_rows[specialty],
    }


def payload_4to(specialty):
    initial_stage = {
        "Mecánica Industrial": "Durante abril, el proyecto se desarrolla en conjunto con Emprendimiento y Empleabilidad para la elaboración de un currículum en inglés y la simulación de entrevistas laborales.",
        "Gráfica": "Durante abril, el proyecto se desarrolla en conjunto con Emprendimiento y Empleabilidad para la elaboración de un currículum en inglés y la simulación de entrevistas laborales.",
        "Electricidad": "Durante abril, el proyecto se desarrolla en conjunto con Emprendimiento y Empleabilidad para la elaboración de un currículum en inglés y la simulación de entrevistas laborales.",
        "Mecánica Automotriz": "A partir de julio, el primer trabajo relevante se centra en la atención al cliente en contexto de taller mecánico, considerando la comunicación profesional con el usuario y la explicación del diagnóstico y del trabajo realizado o por realizar.",
        "Electrónica": "Durante abril, el proyecto se desarrolla en conjunto con Emprendimiento y Empleabilidad para la elaboración de un currículum en inglés y la simulación de entrevistas laborales.",
    }
    reading_focus = {
        "Mecánica Industrial": "Documentos y materiales consultados con especialidad sobre CNC y MasterCAM.",
        "Gráfica": "Documentos y materiales vinculados al equipo recientemente adquirido en la especialidad.",
        "Electricidad": "Textos relacionados con motores eléctricos, magnetismo y terminología clave asociada a estos contenidos.",
        "Mecánica Automotriz": "Textos y pautas de trabajo sobre conceptos de motores y sistemas, mantenimiento del motor, sistemas de frenos, suspensión, dirección y sistema eléctrico, componentes del tren delantero, diagnóstico de fallas y revisiones según distintos hitos de mantenimiento de los vehículos.",
        "Electrónica": "Textos sobre motores eléctricos, automatización considerando sensores, actuadores y PLC, electrónica de potencia y sus usos, servomecanismos y electrónica de consumo.",
    }
    second_eval = {
        "Mecánica Industrial": "Evaluación de comprensión lectora basada en materiales de MasterCAM y CNC.",
        "Gráfica": "Evaluación de comprensión lectora basada en documentación del equipo recientemente adquirido.",
        "Electricidad": "Evaluación de comprensión lectora basada en textos sobre motores eléctricos, magnetismo y terminología técnica clave.",
        "Mecánica Automotriz": "Evaluación de comprensión lectora basada en pautas de trabajo, sistemas automotrices, diagnóstico de fallas y revisiones asociadas a distintos hitos de mantenimiento del vehículo.",
        "Electrónica": "Evaluación de comprensión lectora basada en textos sobre automatización, electrónica de potencia, servomecanismos, motores eléctricos y electrónica de consumo.",
    }
    yearly_scope = {
        "Mecánica Industrial": "Los contenidos técnicos se revisarán progresivamente a lo largo del año mediante comprensión lectora y otros trabajos articulados.",
        "Gráfica": "Los contenidos técnicos se revisarán progresivamente a lo largo del año mediante comprensión lectora y otros trabajos articulados.",
        "Electricidad": "Estos contenidos se revisarán a lo largo del año por medio de comprensión lectora, integrando además otros trabajos de la asignatura.",
        "Mecánica Automotriz": "Estos contenidos se revisarán a lo largo del año por medio de comprensión lectora, integrando además otros trabajos de la asignatura.",
        "Electrónica": "Estos contenidos se revisarán a lo largo del año por medio de comprensión lectora, integrando además otros trabajos de la asignatura.",
    }
    july_task = {
        "Mecánica Industrial": "",
        "Gráfica": "",
        "Electricidad": "",
        "Mecánica Automotriz": "Este trabajo servirá como base para posteriores actividades de comprensión lectora y comunicación técnica durante el segundo semestre.",
        "Electrónica": "",
    }
    first_work = {
        "Mecánica Industrial": "Elaboración de un currículum en inglés y simulación de entrevistas laborales.",
        "Gráfica": "Elaboración de un currículum en inglés y simulación de entrevistas laborales.",
        "Electricidad": "Elaboración de un currículum en inglés y simulación de entrevistas laborales.",
        "Mecánica Automotriz": "Trabajo de atención al cliente en contexto de taller mecánico, centrado en la comunicación con el cliente, la explicación del diagnóstico y la descripción del trabajo a realizar o ya realizado.",
        "Electrónica": "Elaboración de un currículum en inglés y simulación de entrevistas laborales.",
    }
    calendar_rows = {
        "Mecánica Industrial": [
            {"name": "Planificación conjunta con docentes", "months": ["Mar."]},
            {"name": "CV en inglés y entrevista laboral simulada", "months": ["Abr."]},
            {"name": "Lectura guiada de documentos de especialidad", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión lectora", "months": ["Jun.", "Dic."]},
        ],
        "Gráfica": [
            {"name": "Planificación conjunta con docentes", "months": ["Mar."]},
            {"name": "CV en inglés y entrevista laboral simulada", "months": ["Abr."]},
            {"name": "Lectura guiada de documentos de especialidad", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión lectora", "months": ["Jun.", "Dic."]},
        ],
        "Electricidad": [
            {"name": "Planificación conjunta con docentes", "months": ["Mar."]},
            {"name": "CV en inglés y entrevista laboral simulada", "months": ["Abr."]},
            {"name": "Lectura guiada de documentos de especialidad", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión lectora", "months": ["Jun.", "Dic."]},
        ],
        "Mecánica Automotriz": [
            {"name": "Planificación conjunta con docentes", "months": ["Mar.", "Abr.", "May.", "Jun."]},
            {"name": "Trabajo 1: atención al cliente en taller mecánico", "months": ["Jul."]},
            {"name": "Lectura guiada sobre motores, sistemas y diagnóstico de fallas", "months": ["Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión lectora y pautas de trabajo", "months": ["Dic."]},
        ],
        "Electrónica": [
            {"name": "Planificación conjunta con docentes", "months": ["Mar."]},
            {"name": "CV en inglés y entrevista laboral simulada", "months": ["Abr."]},
            {"name": "Lectura guiada de documentos de especialidad", "months": ["May.", "Jun.", "Ago.", "Sep.", "Oct.", "Nov."]},
            {"name": "Evaluación escrita de comprensión lectora", "months": ["Jun.", "Dic."]},
        ],
    }
    return {
        "version_description": f"Creación de planificación de articulación 2026 para 4° Medio - {specialty}.",
        "project_name": f"Inglés para empleabilidad y comprensión técnica en {specialty} - 4° Medio",
        "general_description": (
            f"En 4° Medio, las clases de Inglés mantienen un trabajo articulado con {specialty}, incorporando un foco de empleabilidad y lectura técnica. "
            f"{initial_stage[specialty]} Posteriormente, se profundiza la comprensión lectora a partir de documentos y materiales vinculados a la especialidad. "
            f"{reading_focus[specialty]} {yearly_scope[specialty]} {july_task[specialty]}"
        ),
        "general_objectives": (
            f"Fortalecer la comunicación en Inglés para contextos de empleabilidad y para la comprensión de textos técnicos vinculados a {specialty}, "
            f"articulando la asignatura con los módulos TP y con Emprendimiento y Empleabilidad, y proyectando la revisión de estos contenidos durante el año escolar."
        ),
        "specific_objectives": (
            f"1. Desarrollar como primer trabajo relevante del año: {first_work[specialty]}\n"
            "2. Participar en interacciones orales funcionales en inglés, utilizando estrategias de comunicación pertinentes al contexto técnico o laboral.\n"
            f"3. Comprender documentos y textos técnicos vinculados a {specialty}.\n"
            f"4. Desarrollar evaluaciones y actividades de comprensión lectora en torno a {reading_focus[specialty]}\n"
            f"5. Proyectar la revisión de estos temas a lo largo del año, integrando entre medio otros trabajos de la asignatura. {july_task[specialty]}"
        ),
        "specialties": specialty,
        "subsectors": "Inglés; Emprendimiento y Empleabilidad",
        "expected_results": (
            f"- Desarrollo del primer trabajo del año: {first_work[specialty]}\n"
            "- Interacciones orales funcionales en inglés vinculadas al contexto técnico o laboral.\n"
            f"- Evaluación escrita de comprensión lectora vinculada a {specialty}.\n"
            f"- Desarrollo progresivo de comprensión lectora sobre {reading_focus[specialty]}"
        ),
        "methodology": (
            "Aprendizaje situado, enfoque comunicativo y articulación curricular. Se integra producción escrita funcional, simulación oral de entrevista, "
            f"lectura guiada de documentación técnica y evaluación escrita construida en coordinación con docentes de especialidad. {yearly_scope[specialty]} {july_task[specialty]}"
        ),
        "general_subject": "Inglés",
        "specialty_module": f"Articulación con {specialty} y Emprendimiento y Empleabilidad",
        "articulation_rows": [
            [
                "Empleabilidad y proyección laboral.",
                "Producir textos funcionales y participar en interacciones orales vinculadas al mundo del trabajo.",
                "Demostrar responsabilidad, iniciativa y disposición profesional.",
                f"Relación entre módulos de {specialty}, trayectoria de egreso y preparación para la empleabilidad o la comunicación funcional propia del área.",
                f"Desarrolla el primer trabajo definido para la especialidad: {first_work[specialty]}",
                "A) Comunicarse eficazmente.\nE) Emprender y actuar con iniciativa.",
            ],
            [
                "Simulación de entrevista laboral.",
                "Responder preguntas frecuentes y presentar fortalezas personales en contextos de entrevista en inglés.",
                "Trabajar con seguridad, respeto y escucha activa.",
                f"Uso de experiencia en módulos de {specialty} para responder en una interacción oral funcional, entrevista o atención al cliente según corresponda.",
                "Participa en interacciones orales utilizando vocabulario laboral o técnico y respuestas pertinentes.",
                "C) Resolver situaciones comunicativas en contextos reales.",
            ],
            [
                "Comprensión lectora de documentos técnicos.",
                "Comprender información central y específica en textos técnicos consultados con docentes de especialidad.",
                "Trabajar con autonomía, rigurosidad y pensamiento crítico.",
                reading_focus[specialty],
                second_eval[specialty],
                "D) Gestionar información y utilizarla con criterio.",
            ],
        ],
        "key_action_name": f"Secuencia semestral de empleabilidad y lectura técnica en {specialty}",
        "key_action_objective": f"Articular el Inglés con la vida laboral y con la comprensión de documentación técnica vinculada a {specialty}.",
        "participants": f"Docente de Inglés, docente de Emprendimiento y Empleabilidad, docente(s) de {specialty}, estudiantes de 4° Medio TP y coordinación pedagógica HC-TP.",
        "key_action_description": (
            f"La secuencia incluye como primer trabajo: {first_work[specialty]} Además, contempla interacciones orales funcionales y una evaluación escrita de comprensión "
            f"lectora basada en documentos consultados con docentes de {specialty}. {reading_focus[specialty]} {yearly_scope[specialty]} {july_task[specialty]}"
        ),
        "achievement_indicators": (
            "Redacta un CV funcional en inglés; participa en entrevista simulada usando recursos de apoyo; comprende información explícita e implícita en textos técnicos; "
            f"responde una evaluación escrita pertinente a {specialty}; y progresa en la comprensión de textos asociados a los focos definidos para la especialidad."
        ),
        "verification": "Rúbrica del trabajo inicial, instrumentos para interacción oral funcional, prueba escrita de comprensión lectora y productos de clase.",
        "calendar_rows": calendar_rows[specialty],
    }


def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    build_doc("1ro Medio - Ingles articulacion TP.docx", base_1ro_payload())
    for specialty in SPECIALTIES:
        letter = SPECIALTY_LETTERS[specialty]
        build_doc(f"3ro {letter} - {specialty}.docx", payload_3ro(specialty))
        build_doc(f"4to {letter} - {specialty}.docx", payload_4to(specialty))
    print(f"Generated {len(list(OUTPUT_DIR.glob('*.docx')))} files in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()